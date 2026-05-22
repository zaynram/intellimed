"""
Parse a completed Trust Intake Questionnaire .docx into a TrustData instance.

This is a port of the legacy QuestionnaireParser with the following bug fixes:
- Placeholder stripping uses substring matching instead of exact equality.
- Section detection uses lowercase substring matching with a fallback warning.
- Checkbox parsing scans both paragraphs AND table cells.
- Asset/list-table parsing uses .get() with empty-string defaults.
- Returns a TrustData (not a flat dict).
"""

from __future__ import annotations

import logging
import re
from collections.abc import Generator
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]

from trust_generator.schema import (
    BeneficiaryDeath,
    BeneficiaryShare,
    Child,
    DisputeResolution,
    DistributionStandard,
    DivisionMethod,
    Elections,
    FinancialAccount,
    InitialTrustee,
    InsurancePolicy,
    InsuranceStrategy,
    MarriageInfo,
    OfficeInfo,
    OtherBeneficiary,
    Pension,
    PersonInfo,
    PowerOfAppointment,
    PropertyClassification,
    RealProperty,
    RemoteContingent,
    RetirementStrategy,
    SpecificBequest,
    SuccessorTrustee,
    SurvivingAmendment,
    TangibleDistribution,
    TextBlocks,
    TrustData,
    TrusteeCompensation,
    TrustIdentity,
    TrustType,
    Valuable,
    Vehicle,
    WithdrawalStep,
)

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Placeholder hints — if a cell contains ONLY one of these (possibly with
# surrounding whitespace) it is treated as empty.
# ---------------------------------------------------------------------------

_HINTS: list[str] = [
    "e.g.,",
    "MM/DD/YYYY",
    "XXX-XX-XXXX",
    "Street, City, State, ZIP",
    "If different from Husband",
    "Husband / Wife",
    "Default: Illinois",
    "Yes / No",
    "Describe briefly",
    "Joint / H / W",
    "$",
    "e.g., Checking, IRA, 401k",
    "Trust / POD to ___",
    "e.g., Winnebago",
    "e.g., United States",
    "e.g., John Andrew Doe",
    "e.g., Jane Susan Doe",
    "e.g., The Doe Family Trust",
    "H / W / Both",
    "H / W",
    "Pension / Annuity",
    "e.g., upon college graduation",
    "e.g., 1 year after funding",
    "e.g., 2 years after funding",
    "e.g., 5 years after funding",
    "e.g., 50%",
    "e.g., additional 25%",
    "e.g., remaining 25%",
    "%",
    "H / W / Joint",
]

# ---------------------------------------------------------------------------
# Checkbox mapping  (lowercased phrase → (election field, value))
# ---------------------------------------------------------------------------

_CHECKBOX_MAP: dict[str, tuple[str, str | bool]] = {
    "both husband and wife as co-trustees": ("initial_trustee", "both"),
    "both party a and party b as co-trustees": ("initial_trustee", "both"),
    "husband only": ("initial_trustee", "husband"),
    "party a only": ("initial_trustee", "husband"),
    "wife only": ("initial_trustee", "wife"),
    "party b only": ("initial_trustee", "wife"),
    "all property is communal": ("property_classification", "communal"),
    "some property is separate": ("property_classification", "separate"),
    "equally among all children": ("tangible_distribution", "equal_children"),
    "equally among all beneficiaries": ("tangible_distribution", "equal_beneficiaries"),
    "trustee decides": ("division_method", "trustee"),
    "lottery": ("division_method", "lottery"),
    "sell all and divide": ("division_method", "sell"),
    "hems": ("distribution_standard", "hems"),
    "broader discretion": ("distribution_standard", "broad"),
    "per stirpes to that beneficiary": ("beneficiary_death", "per_stirpes_beneficiary"),
    "per stirpes to our": ("beneficiary_death", "per_stirpes_grantors"),
    "redistribute equally": ("beneficiary_death", "redistribute"),
    "distribute per illinois intestacy": ("remote_contingent", "intestacy"),
    "distribute to a named charity": ("remote_contingent", "charity"),
    "pod/tod directly": ("retirement_strategy", "pod"),
    "payable to the trust (more control": ("retirement_strategy", "trust"),
    "mix of pod/tod and trust": ("retirement_strategy", "mix"),
    "mix (some pod": ("retirement_strategy", "mix"),
    "payable directly to surviving spouse": (
        "insurance_strategy",
        "spouse_then_children",
    ),
    "full power to amend or revoke the entire": ("surviving_amendment", "full"),
    "power to amend only the survivor": ("surviving_amendment", "limited"),
    "trust becomes fully irrevocable": ("surviving_amendment", "irrevocable"),
    "full general power of appointment": ("power_of_appointment", "general"),
    "only among our descendants": ("power_of_appointment", "limited"),
    "assets must pass per the trust": ("power_of_appointment", "none"),
    # Boolean elections — both "Yes" and "No" sides mapped
    "yes (standard": ("no_contest", True),
    "no (we do not want a no-contest": ("no_contest", False),
    "no, do not include a no-contest": ("no_contest", False),
    "yes (strongly recommended)": ("spendthrift", True),
    "no (we do not want spendthrift": ("spendthrift", False),
    "no, do not include spendthrift": ("spendthrift", False),
    "mediation, then arbitration": ("dispute_resolution", "mediation_arbitration"),
    "court proceedings only": ("dispute_resolution", "court"),
    "yes (recommended if any assets": ("probate_coordination", True),
    "no, do not coordinate with probate": ("probate_coordination", False),
    "no (standard": ("trustee_bond", False),
    "yes, require a bond": ("trustee_bond", True),
    "fair and reasonable compensation": ("trustee_compensation", "reasonable"),
    "no compensation for family": ("trustee_compensation", "none"),
    "yes (recommended for most estates)": ("portability", True),
    "no, do not elect portability": ("portability", False),
}

# Checked-box indicators
_CHECKED_PREFIXES: list[str] = [
    "\u2611",  # ☑
    "[X]",
    "[x]",
    "☑",
    "\u2612",  # ☒
]


def _cell_text(cell: object) -> str:
    """Extract text from a docx table cell, stripping placeholder hints.

    Bug fix: uses substring/``in`` matching so that a cell containing
    *only* a hint (possibly with whitespace) is treated as empty, even if
    the exact whitespace or trailing punctuation differs.
    """
    text: str = cell.text.strip()  # type: ignore[union-attr]
    if not text:
        return ""
    text_lower = text.lower()
    for hint in _HINTS:
        if hint.lower() in text_lower:
            # The cell contains a hint — strip it only if there is no
            # other substantive content beyond the hint itself.
            remaining = text_lower.replace(hint.lower(), "", 1).strip()
            # Allow for punctuation-only remnants like trailing commas
            if not remaining or all(c in " \t,;.:-_/" for c in remaining):
                return ""
    return text


def _iterate_table_rows(table: object) -> Generator[list[str], None, None]:
    """Yield cleaned cell text rows, skipping the header row."""
    for row in table.rows[1:]:  # type: ignore[union-attr]
        yield [_cell_text(c) for c in row.cells]  # type: ignore[union-attr]


# ---------------------------------------------------------------------------
# Table section detection
# ---------------------------------------------------------------------------


def _detect_section(table: object, headers: list[str]) -> str | None:
    """Detect which field-answer section a table belongs to.

    Returns the section prefix (husband, wife, marriage, trust_id, office)
    or ``None`` for unknown tables (with a log warning).
    """
    fields_text: list[str] = []
    fi = headers.index("field")
    for row in table.rows[1:]:  # type: ignore[union-attr]
        fields_text.append(_cell_text(row.cells[fi]))
    joined = " ".join(fields_text).lower()

    if "maiden" in joined:
        return "wife"
    if "full legal name" in joined and "date of birth" in joined:
        return "husband"
    if "marriage" in joined or "prenuptial" in joined:
        return "marriage"
    if "trust name" in joined or "governing" in joined:
        return "trust_id"
    if "file" in joined or "attorney assigned" in joined:
        return "office"

    log.warning(
        "Could not detect section for field/answer table; skipping. Fields: %s",
        joined[:120],
    )
    return None


# ---------------------------------------------------------------------------
# Individual table parsers
# ---------------------------------------------------------------------------


def _parse_field_answer(
    table: object, headers: list[str], flat: dict[str, str]
) -> None:
    fi = headers.index("field")
    ai = headers.index("answer")
    section = _detect_section(table, headers)
    if section is None:
        return

    for row in table.rows[1:]:  # type: ignore[union-attr]
        field = _cell_text(row.cells[fi]).strip()
        answer = _cell_text(row.cells[ai]).strip()
        if field and answer:
            key = re.sub(r"[^a-z0-9]+", "_", field.lower()).strip("_")
            flat[f"{section}.{key}"] = answer


def _parse_children(table: object) -> list[dict[str, str]]:
    children: list[dict[str, str]] = []
    for cells in _iterate_table_rows(table):
        if len(cells) > 1 and cells[1].strip():
            children.append({
                "name": cells[1].strip(),
                "dob": cells[2].strip() if len(cells) > 2 else "",
                "relationship": cells[3].strip() if len(cells) > 3 else "",
                "minor": cells[4].strip() if len(cells) > 4 else "",
                "notes": cells[5].strip() if len(cells) > 5 else "",
            })
    return children


def _parse_succession(table: object) -> list[dict[str, str]]:
    items: list[dict[str, str]] = []
    for cells in _iterate_table_rows(table):
        if len(cells) >= 2 and cells[1].strip():
            items.append({
                "order": cells[0].strip(),
                "name": cells[1].strip(),
                "relationship": cells[2].strip() if len(cells) > 2 else "",
                "contact": cells[3].strip() if len(cells) > 3 else "",
            })
    return items


def _parse_list_table(table: object, field_names: list[str]) -> list[dict[str, str]]:
    """Parse a generic list table. Uses .get()-style safe access."""
    items: list[dict[str, str]] = []
    for cells in _iterate_table_rows(table):
        first = cells[0].strip() if cells else ""
        if first and first.upper() not in ("TOTAL:", "TOTAL"):
            item: dict[str, str] = {}
            for i, fname in enumerate(field_names):
                item[fname] = cells[i].strip() if i < len(cells) else ""
            if any(v for v in item.values()):
                items.append(item)
    return items


def _parse_beneficiary_shares(table: object) -> list[dict[str, str]]:
    shares: list[dict[str, str]] = []
    for cells in _iterate_table_rows(table):
        first = cells[0].strip() if cells else ""
        if first and first.upper() not in ("TOTAL:", "TOTAL"):
            shares.append({
                "name": first,
                "relationship": cells[1].strip() if len(cells) > 1 else "",
                "share": cells[2].strip().replace("%", "").strip()
                if len(cells) > 2
                else "",
                "conditions": cells[3].strip() if len(cells) > 3 else "",
            })
    return shares


def _parse_withdrawal(table: object) -> list[dict[str, str]]:
    steps: list[dict[str, str]] = []
    for cells in _iterate_table_rows(table):
        timing = cells[1].strip() if len(cells) > 1 else ""
        percentage = cells[2].strip().replace("%", "").strip() if len(cells) > 2 else ""
        if timing or percentage:
            steps.append({
                "step": cells[0].strip() if cells else "",
                "timing": timing,
                "percentage": percentage,
            })
    return steps


# ---------------------------------------------------------------------------
# Checkbox parsing — scans paragraphs AND table cells
# ---------------------------------------------------------------------------


def _is_checked_line(text: str) -> str | None:
    """If *text* looks like a checked checkbox line, return the cleaned text.

    Returns ``None`` if the line is not a checked checkbox.
    """
    stripped = text.strip()
    if not stripped:
        return None
    for prefix in _CHECKED_PREFIXES:
        if stripped.startswith(prefix):
            clean = stripped[len(prefix) :].strip()
            # Also strip residual bracket chars from patterns like "[X] text"
            clean = clean.lstrip("]").strip()
            return clean
    return None


def _parse_checkboxes(doc: object) -> dict[str, str | bool]:
    """Scan paragraphs AND all table cells for checked checkboxes."""
    elections: dict[str, str | bool] = {}

    # Collect all candidate text lines
    lines: list[str] = []
    for para in doc.paragraphs:  # type: ignore[union-attr]
        lines.append(para.text.strip())
    for table in doc.tables:  # type: ignore[union-attr]
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text.strip())

    for line in lines:
        clean = _is_checked_line(line)
        if clean is None:
            continue
        clean_lower = clean.lower()
        for pattern, (key, val) in _CHECKBOX_MAP.items():
            if pattern in clean_lower:
                elections[key] = val
                break

    return elections


# ---------------------------------------------------------------------------
# Text block parsing
# ---------------------------------------------------------------------------


def _parse_text_blocks(doc: object) -> dict[str, str]:
    markers: dict[str, str] = {
        "in your own words": "statement_of_intent",
        "message below": "personal_message",
        "custom distribution terms": "custom_distribution_terms",
        "custom withdrawal schedule": "custom_beneficiary_terms",
        "anything else you would like": "additional_notes",
    }
    end_markers: list[str] = [
        "section ",
        "article ",
        "for office use",
        "signatures",
        "maps to:",
        "would you like",
        "are there",
        "definition of",
        "method if beneficiaries",
        "by signing below",
    ]

    current: str | None = None
    collected: list[str] = []
    blocks: dict[str, str] = {}

    for para in doc.paragraphs:  # type: ignore[union-attr]
        text = para.text.strip()
        tlow = text.lower()
        for marker, key in markers.items():
            if marker in tlow:
                if current and collected:
                    blocks[current] = "\n".join(collected).strip()
                current = key
                collected = []
                break
        if current:
            is_end = any(tlow.startswith(e) for e in end_markers)
            if is_end and collected:
                blocks[current] = "\n".join(collected).strip()
                current = None
                collected = []
            elif text and not all(c in "_ \t" for c in text):
                collected.append(text)
    if current and collected:
        blocks[current] = "\n".join(collected).strip()

    return blocks


# ---------------------------------------------------------------------------
# Flat dict → TrustData mapping
# ---------------------------------------------------------------------------

# Maps flat key suffixes to PersonInfo field names
_PERSON_KEY_MAP: dict[str, str] = {
    "full_legal_name": "full_legal_name",
    "date_of_birth": "date_of_birth",
    "dob": "date_of_birth",
    "ssn": "ssn",
    "social_security_number": "ssn",
    "address": "address",
    "home_address": "address",
    "phone": "phone",
    "phone_number": "phone",
    "email": "email",
    "email_address": "email",
    "employer": "employer",
    "employer_name": "employer",
    "maiden_name": "maiden_name",
}

_TRUST_ID_KEY_MAP: dict[str, str] = {
    "desired_trust_name": "desired_trust_name",
    "trust_name": "desired_trust_name",
    "date": "date",
    "effective_date": "date",
    "state_of_governing_law": "state_of_governing_law",
    "governing_law_state": "state_of_governing_law",
    "county_of_execution": "county_of_execution",
    "county": "county_of_execution",
    "whose_ssn_for_tax_id": "whose_ssn_for_tax_id",
    "whose_ssn_will_be_used_as_the_trust_tax_id": "whose_ssn_for_tax_id",
}

_MARRIAGE_KEY_MAP: dict[str, str] = {
    "date_of_marriage": "date_of_marriage",
    "state_of_marriage": "state_of_marriage",
    "prenuptial_agreement": "prenuptial_agreement",
    "prenuptial_details": "prenuptial_details",
    "if_yes_describe": "prenuptial_details",
}

_OFFICE_KEY_MAP: dict[str, str] = {
    "file_number": "file_number",
    "attorney": "attorney",
    "attorney_assigned": "attorney",
    "paralegal": "paralegal",
    "paralegal_assigned": "paralegal",
    "date_opened": "date_opened",
}

# Election enum constructors keyed by election field name
_ELECTION_ENUM: dict[str, type] = {
    "initial_trustee": InitialTrustee,
    "property_classification": PropertyClassification,
    "tangible_distribution": TangibleDistribution,
    "division_method": DivisionMethod,
    "distribution_standard": DistributionStandard,
    "beneficiary_death": BeneficiaryDeath,
    "remote_contingent": RemoteContingent,
    "retirement_strategy": RetirementStrategy,
    "insurance_strategy": InsuranceStrategy,
    "surviving_amendment": SurvivingAmendment,
    "power_of_appointment": PowerOfAppointment,
    "dispute_resolution": DisputeResolution,
    "trustee_compensation": TrusteeCompensation,
}


def _map_person(flat: dict[str, str], prefix: str) -> dict[str, str]:
    out: dict[str, str] = {}
    for fk, fv in flat.items():
        if not fk.startswith(prefix + "."):
            continue
        suffix = fk[len(prefix) + 1 :]
        mapped = _PERSON_KEY_MAP.get(suffix)
        if mapped:
            out[mapped] = fv
    return out


def _map_section(
    flat: dict[str, str], prefix: str, key_map: dict[str, str]
) -> dict[str, str]:
    out: dict[str, str] = {}
    for fk, fv in flat.items():
        if not fk.startswith(prefix + "."):
            continue
        suffix = fk[len(prefix) + 1 :]
        mapped = key_map.get(suffix)
        if mapped:
            out[mapped] = fv
    return out


def _build_elections(checkbox_data: dict[str, str | bool]) -> Elections:
    kwargs: dict[str, object] = {}
    for key, val in checkbox_data.items():
        if isinstance(val, bool):
            kwargs[key] = val
        elif key in _ELECTION_ENUM:
            try:
                kwargs[key] = _ELECTION_ENUM[key](val)
            except ValueError:
                log.warning("Unknown election value %s=%r, skipping", key, val)
        else:
            # Might be remote_contingent_charity or similar string field
            kwargs[key] = val
    return Elections(**kwargs)  # type: ignore[arg-type]  # Pydantic coerces at runtime


def _flat_to_trust_data(
    flat: dict[str, str],
    *,
    children: list[dict[str, str]],
    successor_trustees: list[dict[str, str]],
    real_property: list[dict[str, str]],
    financial_accounts: list[dict[str, str]],
    vehicles: list[dict[str, str]],
    insurance_policies: list[dict[str, str]],
    pensions: list[dict[str, str]],
    valuables: list[dict[str, str]],
    beneficiary_shares: list[dict[str, str]],
    specific_bequests: list[dict[str, str]],
    withdrawal_schedule: list[dict[str, str]],
    other_beneficiaries: list[dict[str, str]],
    checkbox_data: dict[str, str | bool],
    text_blocks: dict[str, str],
) -> TrustData:
    """Assemble a TrustData from the various parsed components."""
    party_a_data = _map_person(flat, "husband")
    party_b_data = _map_person(flat, "wife")

    # Auto-detect trust type: if party_a is filled but party_b is empty, individual trust
    party_a_name = party_a_data.get("full_legal_name", "")
    party_b_name = party_b_data.get("full_legal_name", "")

    if bool(party_a_name) ^ bool(party_b_name):
        trust_type = TrustType.INDIVIDUAL
        grantor_data = party_a_data if party_a_name else party_b_data
        log.warning(
            "Auto-detected trust type as INDIVIDUAL because only one grantor name "
            "was provided. If this is a joint trust, ensure both names are filled "
            "in the questionnaire."
        )
    else:
        trust_type = TrustType.JOINT
        grantor_data = {}

    return TrustData(
        trust_type=trust_type,
        grantor=PersonInfo(**grantor_data) if grantor_data else PersonInfo(),
        party_a=PersonInfo(**party_a_data),
        party_b=PersonInfo(**party_b_data),
        marriage=MarriageInfo(**_map_section(flat, "marriage", _MARRIAGE_KEY_MAP)),
        trust_id=TrustIdentity(**_map_section(flat, "trust_id", _TRUST_ID_KEY_MAP)),  # type: ignore[arg-type]  # validator coerces str→SsnOwner
        office=OfficeInfo(**_map_section(flat, "office", _OFFICE_KEY_MAP)),
        children=[Child(**c) for c in children],
        successor_trustees=[SuccessorTrustee(**s) for s in successor_trustees],
        real_property=[RealProperty(**r) for r in real_property],
        financial_accounts=[FinancialAccount(**a) for a in financial_accounts],
        vehicles=[Vehicle(**v) for v in vehicles],
        insurance_policies=[InsurancePolicy(**p) for p in insurance_policies],
        pensions=[Pension(**p) for p in pensions],
        valuables=[Valuable(**v) for v in valuables],
        beneficiary_shares=[BeneficiaryShare(**b) for b in beneficiary_shares],
        specific_bequests=[SpecificBequest(**b) for b in specific_bequests],
        withdrawal_schedule=[WithdrawalStep(**w) for w in withdrawal_schedule],
        other_beneficiaries=[OtherBeneficiary(**o) for o in other_beneficiaries],
        elections=_build_elections(checkbox_data),
        text_blocks=TextBlocks(**text_blocks),
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def parse_docx(filepath: str | Path) -> TrustData:
    """Parse a Trust Intake Questionnaire ``.docx`` file into a TrustData.

    Parameters
    ----------
    filepath:
        Path to the ``.docx`` file.

    Returns
    -------
    TrustData
        Fully populated (with defaults where the questionnaire was blank).
    """
    filepath = Path(filepath)
    log.info("Parsing docx questionnaire: %s", filepath)
    doc = Document(str(filepath))

    # Flat key-value pairs from field/answer tables
    flat: dict[str, str] = {}

    # List-structured data
    children: list[dict[str, str]] = []
    successor_trustees: list[dict[str, str]] = []
    real_property: list[dict[str, str]] = []
    financial_accounts: list[dict[str, str]] = []
    vehicles: list[dict[str, str]] = []
    insurance_policies: list[dict[str, str]] = []
    pensions: list[dict[str, str]] = []
    valuables: list[dict[str, str]] = []
    beneficiary_shares: list[dict[str, str]] = []
    specific_bequests: list[dict[str, str]] = []
    withdrawal_schedule: list[dict[str, str]] = []
    other_beneficiaries: list[dict[str, str]] = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headers = [_cell_text(c).strip().lower() for c in table.rows[0].cells]
        hstr = " ".join(headers)

        if "field" in headers and "answer" in headers:
            _parse_field_answer(table, headers, flat)
        elif "full legal name" in hstr and "dob" in hstr and "#" in headers:
            children = _parse_children(table)
        elif "order" in headers and "full name" in hstr:
            successor_trustees = _parse_succession(table)
        elif "property address" in hstr:
            real_property = _parse_list_table(
                table, ["address", "value", "equity", "transfer"]
            )
        elif "institution" in headers:
            financial_accounts = _parse_list_table(
                table, ["institution", "type", "value", "owner", "designation"]
            )
        elif "year" in hstr and "make" in hstr:
            vehicles = _parse_list_table(
                table, ["description", "vin", "value", "owner", "transfer"]
            )
        elif "company" in headers and "policy" in hstr:
            insurance_policies = _parse_list_table(
                table, ["company", "policy_number", "benefit", "insured", "beneficiary"]
            )
        elif "source" in hstr and "type" in headers:
            pensions = _parse_list_table(
                table, ["source", "type", "value", "owner", "survivor"]
            )
        elif "item description" in hstr:
            valuables = _parse_list_table(
                table, ["description", "value", "owner", "specific_bequest"]
            )
        elif "beneficiary name" in hstr and "share" in hstr:
            beneficiary_shares = _parse_beneficiary_shares(table)
        elif "item" in headers and "recipient" in headers:
            specific_bequests = _parse_list_table(
                table, ["item", "recipient", "instructions"]
            )
        elif "step" in headers and "timing" in headers:
            withdrawal_schedule = _parse_withdrawal(table)
        elif "full name" in hstr and "relationship" in headers and "notes" in hstr:
            other_beneficiaries = _parse_list_table(
                table, ["name", "relationship", "dob", "notes"]
            )
        else:
            log.warning(
                "Unrecognised table (headers: %s) — skipping",
                hstr[:100],
            )

    checkbox_data = _parse_checkboxes(doc)
    text_block_data = _parse_text_blocks(doc)

    td = _flat_to_trust_data(
        flat,
        children=children,
        successor_trustees=successor_trustees,
        real_property=real_property,
        financial_accounts=financial_accounts,
        vehicles=vehicles,
        insurance_policies=insurance_policies,
        pensions=pensions,
        valuables=valuables,
        beneficiary_shares=beneficiary_shares,
        specific_bequests=specific_bequests,
        withdrawal_schedule=withdrawal_schedule,
        other_beneficiaries=other_beneficiaries,
        checkbox_data=checkbox_data,
        text_blocks=text_block_data,
    )

    log.info(
        "Parsed docx successfully — %d flat fields, %d children, %d assets total",
        len(flat),
        len(children),
        len(real_property)
        + len(financial_accounts)
        + len(vehicles)
        + len(insurance_policies)
        + len(pensions)
        + len(valuables),
    )
    return td
