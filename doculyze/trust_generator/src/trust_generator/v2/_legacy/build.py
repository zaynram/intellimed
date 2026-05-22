"""
trust_builder.py
Generates a Family Trust document from parsed questionnaire data.
"""

# ruff: noqa: DTZ005
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.text.font import Font


class TrustGenerator:
    def __init__(self, data):
        self.data = data
        self.doc = Document()
        self._setup()

    def _setup(self):
        style = self.doc.styles["Normal"]

        if isinstance(font := getattr(style, "font", None), Font):
            font.name = "Arial"
            font.size = Pt(11)

        for s in self.doc.sections:
            s.page_width = Inches(8.5)
            s.page_height = Inches(11)
            s.top_margin = Inches(1)
            s.bottom_margin = Inches(1)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)

    # --- Accessors ---
    def g(self, key, default="[NOT PROVIDED]"):
        return self.data.get(key, default)

    @property
    def husband(self):
        return self.g("husband.full_legal_name", "[HUSBAND FULL NAME]")

    @property
    def wife(self):
        return self.g("wife.full_legal_name", "[WIFE FULL NAME]")

    @property
    def trust_name(self):
        return self.g(
            "trust_id.desired_trust_name",
            f"The {self.husband.split()[-1]} Family Trust",
        )

    @property
    def trust_date(self):
        return self.g("trust_id.date", datetime.now().strftime("%B %d, %Y"))

    @property
    def county(self):
        return self.g("trust_id.county_of_execution", "Winnebago")

    @property
    def state(self):
        return self.g("trust_id.state_of_governing_law", "Illinois")

    @property
    def trustee_names(self):
        t = self.g("initial_trustee", "both")
        if t == "both":
            return f"{self.husband} and {self.wife}"
        if t == "husband":
            return self.husband
        if t == "wife":
            return self.wife
        return t

    @property
    def ssn_owner_name(self):
        o = self.g("trust_id.whose_ssn_for_tax_id", "Husband")
        return self.husband if "husband" in o.lower() else self.wife

    # --- Helpers ---
    def h1(self, text):
        h = self.doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.name = "Arial"
            r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    def h2(self, text):
        h = self.doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.name = "Arial"
            r.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

    def body(self, text):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(11)

    def indent(self, text, level=1):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5 * level)
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(11)

    def highlight(self, text):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.highlight_color = 7
        r.bold = True

    def manual_review(self, title, content=""):
        p = self.doc.add_paragraph()
        r = p.add_run(f"*** MANUAL REVIEW: {title} ***")
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.bold = True
        r.font.color.rgb = RGBColor(0xCC, 0, 0)
        r.font.highlight_color = 7

        if content:
            p2 = self.doc.add_paragraph()
            r2 = p2.add_run(content)
            r2.font.name = "Arial"
            r2.font.size = Pt(11)
            r2.font.highlight_color = 6
        else:
            for _ in range(5):
                bl = self.doc.add_paragraph()
                r = bl.add_run(" ")
                r.font.highlight_color = 6

        pe = self.doc.add_paragraph()
        re = pe.add_run("*** END MANUAL REVIEW ***")
        re.font.name = "Arial"
        re.font.size = Pt(10)
        re.bold = True
        re.font.color.rgb = RGBColor(0xCC, 0, 0)
        re.font.highlight_color = 7

    def pb(self):
        self.doc.add_page_break()

    def blank(self, n=1):
        for _ in range(n):
            self.doc.add_paragraph()

    # --- Asset list ---
    def asset_list(self):
        items = []
        for p in self.data.get("real_property", []):
            s = f"Real property at {p['address']}"
            if p.get("equity"):
                s += f" (equity: {p['equity']})"
            items.append(s)
        for a in self.data.get("financial_accounts", []):
            s = f"{a.get('type', 'Account')} at {a['institution']}"
            if a.get("value"):
                s += f" (value: {a['value']})"
            items.append(s)
        for v in self.data.get("vehicles", []):
            s = f"Vehicle: {v['description']}"
            if v.get("value"):
                s += f" (value: {v['value']})"
            items.append(s)
        for p in self.data.get("insurance_policies", []):
            s = f"Life insurance with {p['company']}"
            if p.get("benefit"):
                s += f" (benefit: {p['benefit']})"
            items.append(s)
        for p in self.data.get("pensions", []):
            s = f"{p.get('type', 'Pension')} from {p['source']}"
            if p.get("value"):
                s += f" (value: {p['value']})"
            items.append(s)
        for v in self.data.get("valuables", []):
            s = v["description"]
            if v.get("value"):
                s += f" (value: {v['value']})"
            items.append(s)
        return items or ["[LIST ASSETS]"]

    def withdrawal_text(self):
        steps = self.data.get("withdrawal_schedule", [])
        if not steps:
            return (
                "At intervals to be determined, the beneficiary may withdraw "
                "from the trust share per a staggered schedule. [SPECIFY SCHEDULE]"
            )
        parts = []
        for i, s in enumerate(steps):
            pct = s.get("percentage", "[XX]")
            timing = s.get("timing", "[X years after funding]")
            if i == 0:
                parts.append(
                    f"{pct}% of accumulated trust income and principal, {timing}"
                )
            else:
                parts.append(
                    f"increased by {pct}% of accumulated trust income and "
                    f"principal not already subject to withdrawal, {timing}"
                )
        text = "At the intervals below, the beneficiary may withdraw amounts not exceeding:\n\n"
        for i, p in enumerate(parts):
            text += f"    ({chr(97 + i)}) {p};\n\n"
        text += (
            "These rights are cumulative. A beneficiary exercises this right by written "
            "notice to our Trustee. This right may not be subject to creditor claims or "
            "legal process, and may not be alienated or encumbered."
        )
        return text

    # =================================================================
    # GENERATE
    # =================================================================
    def generate(self, output_path):
        self._title_page()
        self.pb()
        self._article_1()
        self.pb()
        self._article_2()
        self.pb()
        self._article_3()
        self.pb()
        self._article_4()
        self.pb()
        self._article_5()
        self.pb()
        self._article_6()
        self.pb()
        self._article_7()
        self._article_8()
        self.pb()
        self._article_9()
        self.pb()
        self._article_10()
        self.pb()
        self._article_11()
        self.pb()
        self._article_12()
        self.pb()
        self._signatures()
        self.pb()
        self._schedules()
        self.doc.save(output_path)
        return output_path

    # --- Title ---
    def _title_page(self):
        self.blank(4)
        for txt, sz, b in [
            ("THE", 16, True),
            (self.trust_name.upper(), 18, True),
            (f"Dated: {self.trust_date}", 12, False),
        ]:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            r.font.name = "Arial"
            r.font.size = Pt(sz)
            r.bold = b
        self.blank(6)
        for txt in [
            "Prepared by:",
            "Crosby and Crosby LLP",
            "3815 N Mulford Rd. 4",
            "Rockford, IL 61114",
            "(815) 367-6432",
        ]:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            r.font.name = "Arial"
            r.font.size = Pt(11)
            if "Crosby" in txt and "Prepared" not in txt:
                r.bold = True

    # --- Article 1 ---
    def _article_1(self):
        self.h1("Article 1: Establishing Our Trust")
        self.body(
            f"The date of this Trust is {self.trust_date}. The parties are "
            f"{self.husband} and {self.wife} (the \u201cGrantors\u201d) and "
            f"{self.trustee_names} (our \u201cTrustee(s)\u201d). We have transferred "
            f"certain assets to our Trustee(s) to be held in trust subject to this instrument."
        )
        self.body(f"We intend to create a valid trust under the laws of {self.state}.")

        self.h2("1.1 Family")
        self.body(
            f"{self.husband} is \u201cHusband\u201d and {self.wife} is \u201cWife.\u201d "
            f"Both are \u201cwe,\u201d \u201cus,\u201d \u201cour,\u201d or \u201cGrantors.\u201d"
        )
        children = self.data.get("children", [])
        if children:
            self.body(f"We have {len(children)} children:")
            for c in children:
                dob = f", born {c['dob']}" if c.get("dob") else ""
                self.indent(f"{c['name']}{dob}")
            self.body(
                "All references to \u201cour children\u201d mean these children. "
                "\u201cOur descendants\u201d means our children and their descendants."
            )
        else:
            self.highlight("[LIST CHILDREN HERE]")

        other = self.data.get("other_beneficiaries", [])
        if other:
            self.body("We have also provided for:")
            for b in other:
                self.indent(f"{b['name']} ({b.get('relationship', '')})")

        self.h2("1.2 Name of Trust")
        self.body(
            f"This Trust shall be known as the {self.trust_name}. For transfers, "
            f"identify as: \u201c{self.trustee_names}, Trustee(s) of the "
            f"{self.trust_name}, dated {self.trust_date}, and any amendments thereto.\u201d"
        )

        self.h2("1.3 Certificate of Trust and Third Party Reliance")
        self.body(
            "Our Trustee may use a certification of trust instead of this instrument. "
            "Third parties may rely upon it and are not required to inquire into the "
            "Trust\u2019s terms or see to the application of funds received by our Trustee."
        )

        self.h2("1.4 Transferring Property to Our Trust")
        self.body(
            "By executing this instrument, we transfer the property described in the "
            "attached schedules. At the time of drafting, those assets include:"
        )
        for i, a in enumerate(self.asset_list(), 1):
            self.indent(f"{i}. {a}")
        self.body(
            "This instrument also considers additional assets we may obtain before "
            "this instrument becomes irrevocable."
        )
        self.body(
            "Our Trustee accepts and agrees to hold transferred property as Trust "
            "property, to administer for our benefit and our beneficiaries\u2019 benefit."
        )

        pc = self.g("property_classification", "communal")
        if pc == "communal":
            self.body("All property transferred is communal property.")
        else:
            self.body("Separate property is defined in Schedules C and D.")

        self.h2("1.5 Intentions of the Grantors")
        self.body(
            "We create this revocable living trust for our benefit while living, and "
            "for our beneficiaries after our death."
        )
        intent = self.g("statement_of_intent", "")
        if intent and intent != "[NOT PROVIDED]":
            self.manual_review("Statement of Intent (from questionnaire)", intent)
        else:
            self.manual_review("Statement of Intent - Attorney to draft")

        msg = self.g("personal_message", "")
        if msg and msg != "[NOT PROVIDED]":
            self.h2("1.6 Message to Our Beneficiaries")
            self.manual_review("Personal Message (from questionnaire)", msg)

        self.h2("1.7 Grantor Trust Status")
        self.body(
            f"We intend to qualify as a Grantor Trust under IRC Sections 671-677. "
            f"The Tax ID will be {self.ssn_owner_name}\u2019s Social Security number."
        )

        self.h2("1.8 Powers Reserved by Us as Grantors")
        self.body(
            "Either of us may act for the Trust without the other\u2019s consent. "
            "We may jointly amend, restate, or revoke this instrument in writing. "
            "Either may add property; jointly we may remove property. We retain control "
            "of income and principal distributions. Undistributed income is added to principal."
        )

    # --- Article 2 ---
    def _article_2(self):
        self.h1("Article 2: Trustee Succession")

        self.h2("2.1 Resignation")
        self.body(
            "A Trustee may resign by written notice to either of us, or if both "
            "incapacitated or deceased, to beneficiaries and co-trustees."
        )

        self.h2("2.2 General Succession")
        self.body(
            "If either of us ceases to act as Trustee, the surviving spouse continues "
            "as sole Trustee. If both are unable, succession is:"
        )
        for s in self.data.get("successor_trustees", []):
            self.indent(
                f"{s['order']}: {s['name']}"
                + (f" ({s['relationship']})" if s.get("relationship") else "")
            )
        if not self.data.get("successor_trustees"):
            self.highlight("[LIST SUCCESSOR TRUSTEES]")

        self.h2("2.3 Contingent Succession")
        self.body(
            "If no designated Trustee can serve: we may jointly appoint one; if only "
            "one of us survives, that spouse may appoint; if neither survives, a majority "
            "of beneficiaries may appoint or petition the court."
        )

        self.h2("2.4 Removal")
        self.body(
            "We may jointly remove any Trustee. During one\u2019s incapacity, the other "
            "may remove. If both incapacitated, a majority of our children may remove. "
            "After both deaths, a majority of Income Beneficiaries may remove."
        )

        self.h2("2.5 Incapacity of a Trustee")
        self.body(
            "A written declaration of incapacity by a co-Trustee or designated successor, "
            "supported by a physician\u2019s written opinion, terminates the trusteeship. "
            "An objecting Trustee has five days to object and must sign medical releases."
        )

        self.h2("2.6 Rights of Successor Trustees")
        self.body(
            "Successor Trustees have all rights and obligations of the initial Trustee. "
            "No successor need examine prior Trustee\u2019s records."
        )

    # --- Article 3 ---
    def _article_3(self):
        self.h1("Article 3: Administration During Our Lives")

        self.h2("3.1 Right to Amend or Revoke")
        self.body(
            "While both alive, we may revoke or amend in whole or part by written "
            "agreement signed by both of us."
        )

        self.h2("3.2 Surviving Spouse\u2019s Amendment Rights")
        ap = self.g("surviving_amendment", "full")
        if ap == "full":
            self.body(
                "The surviving spouse has full power to amend or revoke the entire Trust."
            )
        elif ap == "limited":
            self.body(
                "The surviving spouse may amend only the Survivor\u2019s Trust and "
                "distributions for predeceased children."
            )
        elif ap == "irrevocable":
            self.body("Upon the first death, this Trust becomes irrevocable.")
        else:
            self.manual_review("Surviving Spouse Amendment Rights")

        self.h2("3.3 Distributions During Our Lives")
        self.body(
            "While both alive, we may use or distribute Trust property as if owned "
            "outside the Trust."
        )

        self.h2("3.4 Administration During Incapacity")
        self.body(
            "If one of us is incapacitated, our Trustee shall distribute for: "
            "(a) the incapacitated Grantor; (b) others for the Grantor\u2019s benefit; "
            "(c) authorized agents; (d) the other Grantor\u2019s HEMS; (e) dependents. "
            "Equal consideration shall be given to both Grantors\u2019 needs."
        )

        self.h2("3.5 Determination of Incapacity")
        self.body(
            "Incapacity is established by two licensed physicians, court declaration, "
            "or unexplained absence exceeding 30 days. Capacity is restored by a "
            "physician\u2019s written opinion."
        )

    # --- Article 4 ---
    def _article_4(self):
        self.h1("Article 4: Administration Upon Death of a Grantor")

        self.h2("4.1 Surviving and Deceased Grantor\u2019s Property")
        self.body(
            "After the first death, the survivor\u2019s interest is the "
            "\u201cSurvivor\u2019s Trust\u201d (Article 5). The deceased\u2019s "
            "interest is the \u201cDeceased Grantor\u2019s Trust Property.\u201d"
        )

        self.h2("4.2 Administrative Trust")
        self.body(
            "The Trust becomes irrevocable as to the deceased\u2019s property. "
            "A separate Tax ID may be needed."
        )

        self.h2("4.3 Payment of Expenses and Taxes")
        self.body(
            "Our Trustee shall pay: (a) funeral and memorial expenses; (b) enforceable "
            "claims; (c) administration expenses; (d) court-ordered allowances; "
            "(e) all taxes including death taxes; (f) any tax-beneficial payments."
        )

        self.h2("4.4 Distribution of Tangible Personal Property")
        td = self.g("tangible_distribution", "equal_children")
        recip = "our children" if td == "equal_children" else "our named beneficiaries"
        self.body(
            f"Our Trustee shall distribute tangible personal property to {recip} in "
            f"substantially equal shares. If unable to agree within six months, the "
            f"Trustee decides."
        )
        for b in self.data.get("specific_bequests", []):
            self.indent(
                f"We direct that {b['item']} be distributed to {b['recipient']}"
                + (f". {b['instructions']}" if b.get("instructions") else ".")
            )

        self.h2("4.5 Tax Elections and Portability")
        self.body(
            "Our Trustee may make necessary tax elections. If the deceased\u2019s "
            "exclusion is not fully used, our Trustee is nominated as executor for "
            "portability under IRC Section 2010(c)(5)(A)."
        )

        if self.g("probate_coordination", str(True)):
            self.h2("4.6 Coordination with Personal Representative")
            self.body(
                "Our Trustee may rely on the Personal Representative\u2019s requests, "
                "accept property from probate, and distribute to the probate estate "
                "when beneficial."
            )

    # --- Article 5 ---
    def _article_5(self):
        self.h1("Article 5: The Survivor\u2019s Trust")
        self.body(
            "Our Trustee allocates the deceased\u2019s remaining property to the "
            "Survivor\u2019s Trust."
        )

        self.h2("5.1 Trustee")
        self.body(
            "The surviving Grantor may serve as sole Trustee and may remove or "
            "replace the Trustee at any time."
        )

        self.h2("5.2 Right to Amend")
        self.body(
            "Subject to Section 3.2, the surviving Grantor may amend, restate, or "
            "revoke the Survivor\u2019s Trust in writing."
        )

        self.h2("5.3 Income and Principal")
        self.body(
            "All net income distributed quarter-annually to the surviving Grantor. "
            "Principal as the surviving Grantor directs for any reason."
        )

        self.h2("5.4 Incapacity")
        self.body("During incapacity, administered per Section 3.4.")

        self.h2("5.5 General Power of Appointment")
        poa = self.g("power_of_appointment", "general")
        if poa == "general":
            self.body(
                "The surviving Grantor may appoint all or any portion at death among "
                "any persons or entities, including creditors of the estate."
            )
        elif poa == "limited":
            self.body("The surviving Grantor may appoint among our descendants only.")
        else:
            self.body("No power of appointment. Property passes per Article 6.")

        self.h2("5.6 After the Surviving Grantor\u2019s Death")
        self.body(
            "The Survivor\u2019s Trust becomes irrevocable. Administered per Article 4, "
            "then the unappointed balance passes per Article 6."
        )

    # --- Article 6 ---
    def _article_6(self):
        self.h1("Article 6: Distribution to Beneficiaries")

        self.h2("6.1 Division of Remaining Property")
        self.body("Our Trustee divides remaining property as follows:")
        shares = self.data.get("beneficiary_shares", [])
        if shares:
            t = self.doc.add_table(rows=len(shares) + 1, cols=4)
            t.style = "Table Grid"
            for i, h in enumerate(["Name", "Relationship", "Share", "Conditions"]):
                c = t.rows[0].cells[i]
                c.text = h
                for r in c.paragraphs[0].runs:
                    r.bold = True
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
            for ri, s in enumerate(shares, 1):
                t.rows[ri].cells[0].text = s.get("name", "")
                t.rows[ri].cells[1].text = s.get("relationship", "")
                t.rows[ri].cells[2].text = f"{s.get('share', '')}%"
                t.rows[ri].cells[3].text = s.get("conditions", "")
                for c in t.rows[ri].cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.name = "Arial"
                            r.font.size = Pt(10)
            self.doc.add_paragraph()
        else:
            self.highlight("[COMPLETE BENEFICIARY TABLE]")

        self.h2("6.2 Standard Distribution Terms")
        std = self.g("distribution_standard", "hems")
        stxt = (
            "health, education, maintenance, or support"
            if std == "hems"
            else "any purpose"
        )
        self.body(
            f"Our Trustee shall distribute income and principal as determined necessary "
            f"for each beneficiary\u2019s {stxt}. Undistributed income added to principal."
        )
        self.body(
            "We request our Trustee be liberal in making distributions, considering "
            "the beneficiary\u2019s other known resources. Principal may be exhausted."
        )
        self.body(self.withdrawal_text())

        self.h2("6.3 If a Beneficiary Is Deceased")
        dp = self.g("beneficiary_death", "per_stirpes_beneficiary")
        if dp == "per_stirpes_beneficiary":
            self.body(
                "The deceased beneficiary\u2019s share passes per stirpes to their "
                "descendants, then to our descendants, then per Article 7."
            )
        elif dp == "per_stirpes_grantors":
            self.body(
                "The share passes per stirpes to our descendants, then per Article 7."
            )
        elif dp == "redistribute":
            self.body(
                "The share is redistributed equally among surviving named beneficiaries."
            )

        custom = self.g("custom_distribution_terms", "")
        custom_b = self.g("custom_beneficiary_terms", "")
        combined = "\n".join(
            filter(
                None,
                [
                    custom if custom != "[NOT PROVIDED]" else "",
                    custom_b if custom_b != "[NOT PROVIDED]" else "",
                ],
            )
        )
        if combined.strip():
            self.h2("6.4 Specific Terms by Beneficiary")
            self.manual_review("Custom terms from questionnaire", combined)

    # --- Article 7 ---
    def _article_7(self):
        self.h1("Article 7: Remote Contingent Distribution")
        rc = self.g("remote_contingent", "intestacy")
        if rc == "intestacy":
            self.body(
                f"If no one is qualified to receive distribution, property passes "
                f"per {self.state} intestacy laws."
            )
        elif rc == "charity":
            self.body(
                "If no one is qualified, property passes to "
                f"{self.g('remote_contingent_charity', '[NAMED CHARITY]')}."
            )

    # --- Article 8 ---
    def _article_8(self):
        self.h1("Article 8: Underage and Incapacitated Beneficiaries")
        self.body(
            "For beneficiaries under 18 or incapacitated, our Trustee may distribute "
            "directly, to a guardian, as UTMA custodian, or to an agent. Alternatively, "
            "property may be retained in a separate trust for HEMS distributions. Upon "
            "reaching 18 or restoration of capacity, the beneficiary may withdraw all "
            "accumulated income and principal."
        )

    # --- Article 9 ---
    def _article_9(self):
        self.h1("Article 9: Retirement Plans and Life Insurance")

        self.h2("9.1 Retirement Plans")
        self.body(
            "Our Trustee may determine manner and timing of retirement plan payments "
            "consistent with IRC Section 401(a)(9). Our Trustee may not change plan "
            "beneficiaries."
        )
        rs = self.g("retirement_strategy", "pod")
        if rs == "trust":
            self.manual_review(
                "Retirement Plan Trust Provisions",
                "Client elected Trust-controlled retirement. Draft conduit/see-through "
                "provisions and review SECURE Act compliance.",
            )
        elif rs == "mix":
            self.manual_review(
                "Mixed Retirement Strategy",
                "Review Section 4 asset table for per-account designations.",
            )

        self.h2("9.2 Life Insurance")
        self.body(
            "During our lives, each reserves all rights over insurance policies. "
            "After death, our Trustee collects proceeds payable to the Trust. "
            "Proceeds must never become part of our probate estate."
        )

    # --- Article 10 ---
    def _article_10(self):
        self.h1("Article 10: Trust Administration")

        self.h2("10.1 Distributions")
        self.body(
            "Our Trustee may make cash, in-kind, or mixed distributions in any "
            "proportion without regard to tax attributes or beneficiary consent."
        )

        self.h2("10.2 No Court Proceedings; No Bond")
        self.body(
            "Our Trustee shall administer with efficiency and freedom from judicial "
            "intervention. No bond is required."
        )
        if self.g("dispute_resolution") == "mediation_arbitration":
            self.body(
                "Disputes should be resolved by mediation, then arbitration per the "
                "Uniform Arbitration Act."
            )

        self.h2("10.3 Trustee Compensation")
        comp = self.g("trustee_compensation", "reasonable")
        if comp == "reasonable":
            self.body(
                "Individual Trustees receive fair and reasonable compensation. "
                "Corporate fiduciaries per their fee schedule. All may be reimbursed."
            )
        elif comp == "none":
            self.body(
                "Family Trustees serve without compensation. Corporate fiduciaries "
                "per their fee schedule. All may be reimbursed."
            )

        self.h2("10.4 Accounting")
        self.body(
            "After the first death, annual accounting to Income Beneficiaries unless "
            "waived. Objections within 60 days. Records available at reasonable times."
        )

        self.h2("10.5 Liability Limitations")
        self.body(
            "No Trustee liability unless bad faith shown by clear and convincing evidence."
        )

        self.h2("10.6 Merge, Sever, or Terminate Trusts")
        self.body(
            "Our Trustee may merge, sever, or terminate trusts when economical. "
            "Terminated trust property distributed to: us if living; surviving Grantor; "
            "mandatory income beneficiaries; then discretionary income beneficiaries."
        )

    # --- Article 11 ---
    def _article_11(self):
        self.h1("Article 11: Trustee\u2019s Powers")
        self.body(
            f"Our Trustee has all powers granted by this Trust and {self.state} law, "
            f"including the Illinois Trust Code (760 ILCS 3/101, et seq.), incorporated "
            f"herein. Powers include: invest, bank, conduct business, sell/transfer, "
            f"contract, manage real estate, acquire residences, retain property, "
            f"borrow, hold in nominee form, insure, settle claims, execute documents."
        )

        self.h2("11.1 Limitations")
        self.body(
            "An Interested Trustee is limited to HEMS distributions per IRC Sections "
            "2041 and 2514. When prohibited from acting, an Independent Special Trustee "
            "may be appointed. No Trustee may discharge their own legal obligations. "
            "These limitations do not apply while either of us serves as Trustee of "
            "the Survivor\u2019s Trust."
        )

    # --- Article 12 ---
    def _article_12(self):
        self.h1("Article 12: General Provisions")

        self.h2("12.1 Maximum Term")
        self.body(
            "Each trust terminates 21 years after the death of the last descendant of "
            "our grandparents living at the time of the first of us to die."
        )

        if self.g("spendthrift", str(True)):
            self.h2("12.2 Spendthrift Provision")
            self.body(
                "All trusts are spendthrift trusts. No beneficiary may assign, "
                "anticipate, or encumber any interest. Neither income nor principal "
                "is subject to creditors or involuntary transfer."
            )

        if self.g("no_contest", str(True)):
            self.h2("12.3 Contest Provision")
            self.body(
                "Any person who contests this Trust forfeits their share and is "
                "deemed to have predeceased the last of us."
            )

        self.h2("12.4 Survivorship Presumption")
        self.body(
            "If we die simultaneously, each is deemed to have predeceased the other. "
            "A beneficiary dying within 30 days of a Grantor is deemed to have "
            "predeceased the Grantor."
        )

        self.h2("12.5 Definitions")
        defs = [
            (
                "Adopted Persons",
                "Legally adopted before 18 have same rights as natural children.",
            ),
            ("Descendants", "Children, grandchildren, etc. Not nieces or nephews."),
            (
                "Education",
                "An ascertainable standard: schooling, college, vocational, tuition, room, board, fees.",
            ),
            ("Grantor", "Same as Settlor or Trustor."),
            ("Incapacity", "As defined in Section 3.5."),
            ("Income Beneficiary", "Beneficiary entitled to net income distributions."),
            (
                "Independent Trustee",
                "Not related or subordinate per IRC Section 672(c).",
            ),
            (
                "Per Stirpes",
                "Divided among children and deceased children\u2019s descendants equally.",
            ),
        ]
        for term, defn in defs:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            rb = p.add_run(f"{term}: ")
            rb.bold = True
            rb.font.name = "Arial"
            rb.font.size = Pt(11)
            rn = p.add_run(defn)
            rn.font.name = "Arial"
            rn.font.size = Pt(11)

        self.h2("12.6 Rules of Construction")
        self.body(
            f"Executed in counterparts. Governed by {self.state} law. Headings for "
            f"convenience. Notices in writing. Invalid provisions do not affect others. "
            f"No antilapse statutes apply."
        )

    # --- Signatures ---
    def _signatures(self):
        self.h1("Execution")
        self.body(f"Executed on {self.trust_date}. Effective when signed by us.")
        self.blank(3)
        self.body("________________________________________")
        self.body(f"{self.husband}, Grantor and Trustee")
        self.blank(2)
        self.body("________________________________________")
        self.body(f"{self.wife}, Grantor and Trustee")
        self.blank(2)
        self.body(f"STATE OF {self.state.upper()} )")
        self.body(") ss.")
        self.body(f"COUNTY OF {self.county.upper()} )")
        self.blank()
        self.body(
            f"Acknowledged before me on {self.trust_date}, by {self.husband}, "
            f"as Grantor and Trustee, and {self.wife}, as Grantor and Trustee."
        )
        self.blank(2)
        self.body("[Seal]")
        self.blank()
        self.body("________________________________________")
        self.body("Notary Public")
        self.body("My commission expires: _______________")

    # --- Schedules ---
    def _schedules(self):
        self.h1("Schedule A: Communal Property")
        self.body("Transferred to this Trust:")
        self.blank()
        self.body("Ten Dollars Cash")
        self.manual_review("Additional communal property")

        self.pb()
        self.h1("Schedule B: Memorandum of Distribution [OPTIONAL]")
        bq = self.data.get("specific_bequests", [])
        if bq:
            for b in bq:
                self.indent(
                    f"{b['item']} \u2192 {b['recipient']}"
                    + (f" ({b['instructions']})" if b.get("instructions") else "")
                )
        else:
            self.manual_review("Specific bequests if applicable")

        if self.g("property_classification") == "separate":
            self.pb()
            self.h1("Schedule C: Husband\u2019s Separate Property")
            self.manual_review("Husband\u2019s separate property")
            self.pb()
            self.h1("Schedule D: Wife\u2019s Separate Property")
            self.manual_review("Wife\u2019s separate property")
