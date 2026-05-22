"""
questionnaire_parser.py
Parses a completed Trust Intake Questionnaire .docx and extracts all answers.
"""

import re

from docx import Document


class QuestionnaireParser:
    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.data = {}

    def parse(self):
        self._parse_tables()
        self._parse_checkboxes()
        self._parse_text_blocks()
        self._apply_defaults()
        return self.data

    def _cell_text(self, cell):
        text = cell.text.strip()
        hints = [
            "e.g.,",
            "MM/DD/YYYY",
            "XXX-XX-XXXX",
            "Street, City, State, ZIP",
            "If different from Husband",
            "Husband / Wife",
            "Default: Illinois",
            "Yes / No",
            "Describe briefly",
            "Joint / H / W",
            "$",
            "e.g., Checking, IRA, 401k",
            "Trust / POD to ___",
            "e.g., Winnebago",
            "e.g., United States",
            "e.g., John Andrew Doe",
            "e.g., Jane Susan Doe",
            "e.g., The Doe Family Trust",
            "H / W / Both",
            "H / W",
            "Pension / Annuity",
            "e.g., upon college graduation",
            "e.g., 1 year after funding",
            "e.g., 2 years after funding",
            "e.g., 5 years after funding",
            "e.g., 50%",
            "e.g., additional 25%",
            "e.g., remaining 25%",
            "%",
            "H / W / Joint",
            "Yes / No",
        ]
        for hint in hints:
            if text == hint:
                return ""
        return text

    def _parse_tables(self):
        for table in self.doc.tables:
            if len(table.rows) < 2:
                continue
            headers = [self._cell_text(c).strip().lower() for c in table.rows[0].cells]
            hstr = " ".join(headers)

            if "field" in headers and "answer" in headers:
                self._parse_field_answer(table, headers)
            elif "full legal name" in hstr and "dob" in hstr and "#" in headers:
                self._parse_children(table)
            elif "order" in headers and "full name" in hstr:
                self._parse_succession(table)
            elif "property address" in hstr:
                self._parse_list_table(
                    table, "real_property", ["address", "value", "equity", "transfer"]
                )
            elif "institution" in headers:
                self._parse_list_table(
                    table,
                    "financial_accounts",
                    ["institution", "type", "value", "owner", "designation"],
                )
            elif "year" in hstr and "make" in hstr:
                self._parse_list_table(
                    table,
                    "vehicles",
                    ["description", "vin", "value", "owner", "transfer"],
                )
            elif "company" in headers and "policy" in hstr:
                self._parse_list_table(
                    table,
                    "insurance_policies",
                    ["company", "policy_number", "benefit", "insured", "beneficiary"],
                )
            elif "source" in hstr and "type" in headers:
                self._parse_list_table(
                    table, "pensions", ["source", "type", "value", "owner", "survivor"]
                )
            elif "item description" in hstr:
                self._parse_list_table(
                    table,
                    "valuables",
                    ["description", "value", "owner", "specific_bequest"],
                )
            elif "beneficiary name" in hstr and "share" in hstr:
                self._parse_beneficiary_shares(table)
            elif "item" in headers and "recipient" in headers:
                self._parse_list_table(
                    table, "specific_bequests", ["item", "recipient", "instructions"]
                )
            elif "step" in headers and "timing" in headers:
                self._parse_withdrawal(table)
            elif "full name" in hstr and "relationship" in headers and "notes" in hstr:
                self._parse_list_table(
                    table,
                    "other_beneficiaries",
                    ["name", "relationship", "dob", "notes"],
                )

    def _parse_field_answer(self, table, headers):
        fi = headers.index("field")
        ai = headers.index("answer")
        fields = [self._cell_text(r.cells[fi]) for r in table.rows[1:]]
        joined = " ".join(fields).lower()
        if "maiden" in joined:
            section = "wife"
        elif "full legal name" in joined and "date of birth" in joined:
            section = "husband"
        elif "marriage" in joined or "prenuptial" in joined:
            section = "marriage"
        elif "trust name" in joined or "governing" in joined:
            section = "trust_id"
        elif "file" in joined or "attorney assigned" in joined:
            section = "office"
        else:
            section = "unknown"

        for row in table.rows[1:]:
            field = self._cell_text(row.cells[fi]).strip()
            answer = self._cell_text(row.cells[ai]).strip()
            if field and answer:
                key = re.sub(r"[^a-z0-9]+", "_", field.lower()).strip("_")
                self.data[f"{section}.{key}"] = answer

    def _parse_children(self, table):
        children = []
        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if len(cells) > 1 and cells[1].strip():
                children.append({
                    "name": cells[1].strip(),
                    "dob": cells[2].strip() if len(cells) > 2 else "",
                    "relationship": cells[3].strip() if len(cells) > 3 else "",
                    "minor": cells[4].strip() if len(cells) > 4 else "",
                    "notes": cells[5].strip() if len(cells) > 5 else "",
                })
        self.data["children"] = children

    def _parse_succession(self, table):
        items = []
        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if len(cells) >= 2 and cells[1].strip():
                items.append({
                    "order": cells[0].strip(),
                    "name": cells[1].strip(),
                    "relationship": cells[2].strip() if len(cells) > 2 else "",
                    "contact": cells[3].strip() if len(cells) > 3 else "",
                })
        self.data["successor_trustees"] = items

    def _parse_list_table(self, table, key, field_names):
        items = []
        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if cells[0].strip() and cells[0].strip().upper() != "TOTAL:":
                item = {}
                for i, fname in enumerate(field_names):
                    item[fname] = cells[i].strip() if i < len(cells) else ""
                if any(v for v in item.values()):
                    items.append(item)
        self.data[key] = items

    def _parse_beneficiary_shares(self, table):
        shares = []
        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if cells[0].strip() and cells[0].strip().upper() not in ("TOTAL:", "TOTAL"):
                shares.append({
                    "name": cells[0].strip(),
                    "relationship": cells[1].strip() if len(cells) > 1 else "",
                    "share": cells[2].strip().replace("%", "").strip()
                    if len(cells) > 2
                    else "",
                    "conditions": cells[3].strip() if len(cells) > 3 else "",
                })
        self.data["beneficiary_shares"] = shares

    def _parse_withdrawal(self, table):
        steps = []
        for row in table.rows[1:]:
            cells = [self._cell_text(c) for c in row.cells]
            if cells[1].strip() or (len(cells) > 2 and cells[2].strip()):
                steps.append({
                    "step": cells[0].strip(),
                    "timing": cells[1].strip(),
                    "percentage": cells[2].strip().replace("%", "").strip()
                    if len(cells) > 2
                    else "",
                })
        self.data["withdrawal_schedule"] = steps

    def _parse_checkboxes(self):
        checked = "\u2611"
        mapping = {
            "both husband and wife as co-trustees": ("initial_trustee", "both"),
            "husband only": ("initial_trustee", "husband"),
            "wife only": ("initial_trustee", "wife"),
            "all property is communal": ("property_classification", "communal"),
            "some property is separate": ("property_classification", "separate"),
            "equally among all children": ("tangible_distribution", "equal_children"),
            "equally among all beneficiaries": (
                "tangible_distribution",
                "equal_beneficiaries",
            ),
            "trustee decides": ("division_method", "trustee"),
            "lottery": ("division_method", "lottery"),
            "sell all and divide": ("division_method", "sell"),
            "hems": ("distribution_standard", "hems"),
            "broader discretion": ("distribution_standard", "broad"),
            "per stirpes to that beneficiary": (
                "beneficiary_death",
                "per_stirpes_beneficiary",
            ),
            "per stirpes to our": ("beneficiary_death", "per_stirpes_grantors"),
            "redistribute equally": ("beneficiary_death", "redistribute"),
            "distribute per illinois intestacy": ("remote_contingent", "intestacy"),
            "distribute to a named charity": ("remote_contingent", "charity"),
            "pod/tod directly": ("retirement_strategy", "pod"),
            "payable to the trust (more control": ("retirement_strategy", "trust"),
            "mix": ("retirement_strategy", "mix"),
            "payable directly to surviving spouse": (
                "insurance_strategy",
                "spouse_then_children",
            ),
            "full power to amend or revoke the entire": ("surviving_amendment", "full"),
            "power to amend only the survivor": ("surviving_amendment", "limited"),
            "trust becomes fully irrevocable": ("surviving_amendment", "irrevocable"),
            "full general power of appointment": ("power_of_appointment", "general"),
            "only among our descendants": ("power_of_appointment", "limited"),
            "assets must pass per the trust": ("power_of_appointment", "none"),
            "yes (standard": ("no_contest", True),
            "yes (strongly recommended)": ("spendthrift", True),
            "mediation, then arbitration": (
                "dispute_resolution",
                "mediation_arbitration",
            ),
            "court proceedings only": ("dispute_resolution", "court"),
            "yes (recommended if any assets": ("probate_coordination", True),
            "no (standard": ("trustee_bond", False),
            "fair and reasonable compensation": ("trustee_compensation", "reasonable"),
            "no compensation for family": ("trustee_compensation", "none"),
            "yes (recommended for most estates)": ("portability", True),
        }

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_checked = text.startswith((checked, "[X]", "[x]"))
            if is_checked:
                clean = text.lstrip(checked).lstrip("[Xx]").strip()
                for pattern, (key, val) in mapping.items():
                    if pattern in clean.lower():
                        self.data[key] = val
                        break

    def _parse_text_blocks(self):
        markers = {
            "in your own words": "statement_of_intent",
            "message below": "personal_message",
            "custom distribution terms": "custom_distribution_terms",
            "custom withdrawal schedule": "custom_beneficiary_terms",
            "anything else you would like": "additional_notes",
        }
        end_markers = [
            "section ",
            "article ",
            "for office use",
            "signatures",
            "maps to:",
            "would you like",
            "are there",
            "definition of",
            "method if beneficiaries",
            "by signing below",
        ]

        current = None
        collected = []
        for para in self.doc.paragraphs:
            text = para.text.strip()
            tlow = text.lower()
            for marker, key in markers.items():
                if marker in tlow:
                    if current and collected:
                        self.data[current] = "\n".join(collected).strip()
                    current = key
                    collected = []
                    break
            if current:
                is_end = any(tlow.startswith(e) for e in end_markers)
                if is_end and collected:
                    self.data[current] = "\n".join(collected).strip()
                    current = None
                    collected = []
                elif text and not all(c in "_ \t" for c in text):
                    collected.append(text)
        if current and collected:
            self.data[current] = "\n".join(collected).strip()

    def _apply_defaults(self):
        defaults = {
            "trust_id.state_of_governing_law": "Illinois",
            "trust_id.county_of_execution": "Winnebago",
            "initial_trustee": "both",
            "property_classification": "communal",
            "tangible_distribution": "equal_children",
            "division_method": "trustee",
            "distribution_standard": "hems",
            "beneficiary_death": "per_stirpes_beneficiary",
            "remote_contingent": "intestacy",
            "retirement_strategy": "pod",
            "insurance_strategy": "spouse_then_children",
            "portability": True,
            "surviving_amendment": "full",
            "power_of_appointment": "general",
            "no_contest": True,
            "spendthrift": True,
            "dispute_resolution": "mediation_arbitration",
            "probate_coordination": True,
            "trustee_bond": False,
            "trustee_compensation": "reasonable",
        }
        for k, v in defaults.items():
            if k not in self.data:
                self.data[k] = v
