"""
Reusable .docx formatting helpers.

Extracts the formatting methods from the legacy TrustGenerator so they can be
shared across generators (trust document, clean printable questionnaire, etc.).
The visual output matches the original exactly: Arial 11pt, justified body,
dark-blue headings, red/yellow manual review blocks.
"""

from __future__ import annotations

from docx import Document as _make_document
from docx.document import Document
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_COLOR_INDEX,
)
from docx.shared import Inches, Pt, RGBColor
from docx.text.font import Font


class DocxFormatter:
    """Stateful .docx formatter wrapping a python-docx Document."""

    def __init__(self, doc: Document | None = None) -> None:
        self.doc: Document = doc or _make_document()
        self._setup()

    def _setup(self) -> None:
        style = self.doc.styles["Normal"]
        if isinstance(font := getattr(style, "font", None), Font):
            font.name = "Arial"
            font.size = Pt(11)

        for s in self.doc.sections:
            s.page_width = Inches(8.5)
            s.page_height = Inches(11)
            s.top_margin = Inches(1)
            s.bottom_margin = Inches(1)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)

    def h1(self, text: str) -> None:
        h = self.doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.name = "Arial"
            r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    def h2(self, text: str) -> None:
        h = self.doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.name = "Arial"
            r.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

    def body(self, text: str) -> None:
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(11)

    def indent(self, text: str, level: int = 1) -> None:
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5 * level)
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(11)

    def highlight(self, text: str) -> None:
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        r.bold = True

    def manual_review(self, title: str, content: str = "") -> None:
        p = self.doc.add_paragraph()
        r = p.add_run(f"*** MANUAL REVIEW: {title} ***")
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.bold = True
        r.font.color.rgb = RGBColor(0xCC, 0, 0)
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW

        if content:
            p2 = self.doc.add_paragraph()
            r2 = p2.add_run(content)
            r2.font.name = "Arial"
            r2.font.size = Pt(11)
            r2.font.highlight_color = WD_COLOR_INDEX.RED
        else:
            for _ in range(5):
                bl = self.doc.add_paragraph()
                r_blank = bl.add_run(" ")
                r_blank.font.highlight_color = WD_COLOR_INDEX.RED

        pe = self.doc.add_paragraph()
        re = pe.add_run("*** END MANUAL REVIEW ***")
        re.font.name = "Arial"
        re.font.size = Pt(10)
        re.bold = True
        re.font.color.rgb = RGBColor(0xCC, 0, 0)
        re.font.highlight_color = WD_COLOR_INDEX.YELLOW

    def pb(self) -> None:
        self.doc.add_page_break()

    def blank(self, n: int = 1) -> None:
        for _ in range(n):
            self.doc.add_paragraph()

    def centered(self, text: str, size: int = 11, bold: bool = False) -> None:
        """Add a centered paragraph with custom font size."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.bold = bold

    def definition(self, term: str, text: str) -> None:
        """Add a bold-term + normal-definition paragraph."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        rb = p.add_run(f"{term}: ")
        rb.bold = True
        rb.font.name = "Arial"
        rb.font.size = Pt(11)
        rn = p.add_run(text)
        rn.font.name = "Arial"
        rn.font.size = Pt(11)

    def add_table(
        self,
        headers: list[str],
        rows: list[list[str]],
    ) -> None:
        """Add a formatted table with headers and data rows."""
        t = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        t.style = "Table Grid"
        for i, header in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = header
            for r in cell.paragraphs[0].runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
        for ri, row_data in enumerate(rows, 1):
            for ci, value in enumerate(row_data):
                t.rows[ri].cells[ci].text = value
                for p in t.rows[ri].cells[ci].paragraphs:
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(10)
        self.doc.add_paragraph()

    def save(self, path: str) -> str:
        self.doc.save(path)
        return path
