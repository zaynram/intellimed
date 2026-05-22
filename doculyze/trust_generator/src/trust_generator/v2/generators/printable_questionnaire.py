"""
Generate a clean, printable Trust Intake Questionnaire .docx.

The output contains NO placeholder or hint text — every answer cell is blank
and ready for handwriting or typing.  Section structure mirrors the canonical
TrustData schema so that every field the parser expects is present.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx.enum.table import WD_ROW_HEIGHT_RULE  # type: ignore[import-untyped]
from docx.shared import Inches  # type: ignore[import-untyped]

from trust_generator.v2.config import AppConfig, load_config
from trust_generator.v2.generators.docx_formatter import DocxFormatter

log = logging.getLogger(__name__)

# Minimum row height (in inches) so handwritten answers fit comfortably.
_ROW_HEIGHT = Inches(0.4)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _blank_rows(n: int, cols: int) -> list[list[str]]:
    """Return *n* rows of *cols* empty strings."""
    return [[""] * cols for _ in range(n)]


def _set_row_heights(fmt: DocxFormatter, *, start: int = 1) -> None:
    """Set minimum row height on the most-recently-added table."""
    table = fmt.doc.tables[-1]
    for row in table.rows[start:]:
        row.height = _ROW_HEIGHT
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _add_blank_table(
    fmt: DocxFormatter,
    headers: list[str],
    blank_count: int,
) -> None:
    """Convenience: add a table with *blank_count* empty data rows."""
    fmt.add_table(headers, _blank_rows(blank_count, len(headers)))
    _set_row_heights(fmt)


# ---------------------------------------------------------------------------
# Checkbox helpers
# ---------------------------------------------------------------------------

_EMPTY_BOX = "\u2610"  # ☐


def _checkbox_group(fmt: DocxFormatter, title: str, options: list[str]) -> None:
    """Render a group of empty checkboxes with a bold label."""
    fmt.body(f"{title}:")
    for opt in options:
        fmt.indent(f"{_EMPTY_BOX}  {opt}")


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def generate_printable_questionnaire(
    output_path: str | Path,
    config: AppConfig | None = None,
    trust_type: str = "joint",
    *,
    party_a_label: str = "Husband",
    party_b_label: str = "Wife",
) -> str:
    """Generate a clean, blank Trust Intake Questionnaire .docx.

    Returns the output path as a string.
    """
    cfg = config or load_config()
    fmt = DocxFormatter()

    _header(fmt, cfg)
    _section_office(fmt)
    if trust_type == "individual":
        _section_grantor(fmt)
    else:
        _section_party_a(fmt, party_a_label)
        _section_party_b(fmt, party_b_label)
        _section_marriage(fmt)
    _section_trust_info(fmt)
    fmt.pb()
    _section_children(fmt)
    _section_successor_trustees(fmt)
    fmt.pb()
    _section_real_property(fmt)
    _section_financial_accounts(fmt)
    fmt.pb()
    _section_vehicles(fmt)
    _section_insurance(fmt)
    _section_pensions(fmt)
    fmt.pb()
    _section_valuables(fmt)
    _section_beneficiary_shares(fmt)
    _section_specific_bequests(fmt)
    fmt.pb()
    _section_withdrawal_schedule(fmt)
    _section_elections(fmt, trust_type=trust_type)
    fmt.pb()
    _section_statement_of_intent(fmt)
    _section_personal_message(fmt)
    _section_additional_notes(fmt)
    _section_signature(fmt)

    out = str(Path(output_path))
    fmt.save(out)
    log.info("Printable questionnaire written to %s", out)
    return out


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------


def _header(fmt: DocxFormatter, cfg: AppConfig) -> None:
    fmt.centered(cfg.firm.name, size=16, bold=True)
    fmt.centered(cfg.firm.address_line1, size=10)
    fmt.centered(cfg.firm.address_line2, size=10)
    fmt.centered(cfg.firm.phone, size=10)
    fmt.blank()
    fmt.centered("Trust Intake Questionnaire", size=14, bold=True)
    fmt.centered(f"Date: {'_' * 40}", size=11)
    fmt.blank()


def _section_office(fmt: DocxFormatter) -> None:
    fmt.h2("Office Use")
    fmt.add_table(
        ["Field", ""],
        [
            ["File Number", ""],
            ["Attorney", ""],
            ["Paralegal", ""],
            ["Date Opened", ""],
        ],
    )
    _set_row_heights(fmt)


def _person_table(fmt: DocxFormatter, fields: list[str]) -> None:
    fmt.add_table(["Field", ""], [[f, ""] for f in fields])
    _set_row_heights(fmt)


def _section_grantor(fmt: DocxFormatter) -> None:
    fmt.h2("Grantor Information")
    _person_table(
        fmt,
        [
            "Full Legal Name",
            "Date of Birth",
            "Social Security Number",
            "Address",
            "Phone",
            "Email",
            "Employer",
        ],
    )


def _section_party_a(fmt: DocxFormatter, label: str = "Husband") -> None:
    fmt.h2(f"{label} Information")
    _person_table(
        fmt,
        [
            "Full Legal Name",
            "Date of Birth",
            "Social Security Number",
            "Address",
            "Phone",
            "Email",
            "Employer",
        ],
    )


def _section_party_b(fmt: DocxFormatter, label: str = "Wife") -> None:
    fmt.h2(f"{label} Information")
    _person_table(
        fmt,
        [
            "Full Legal Name",
            "Date of Birth",
            "Social Security Number",
            "Address",
            "Phone",
            "Email",
            "Employer",
            "Maiden Name",
        ],
    )


def _section_marriage(fmt: DocxFormatter) -> None:
    fmt.h2("Marriage Information")
    fmt.add_table(
        ["Field", ""],
        [
            ["Date of Marriage", ""],
            ["State of Marriage", ""],
            ["Prenuptial Agreement", ""],
            ["Prenuptial Agreement Details", ""],
        ],
    )
    _set_row_heights(fmt)


def _section_trust_info(fmt: DocxFormatter) -> None:
    fmt.h2("Trust Information")
    fmt.add_table(
        ["Field", ""],
        [
            ["Desired Trust Name", ""],
            ["Trust Date", ""],
            ["State of Governing Law", ""],
            ["County of Execution", ""],
            ["Whose SSN for Tax ID", ""],
        ],
    )
    _set_row_heights(fmt)


def _section_children(fmt: DocxFormatter) -> None:
    fmt.h2("Children")
    _add_blank_table(
        fmt,
        ["#", "Full Legal Name", "Date of Birth", "Relationship", "Minor?", "Notes"],
        6,
    )


def _section_successor_trustees(fmt: DocxFormatter) -> None:
    fmt.h2("Successor Trustees")
    _add_blank_table(
        fmt,
        ["Order", "Full Name", "Relationship", "Contact Info"],
        4,
    )


def _section_real_property(fmt: DocxFormatter) -> None:
    fmt.h2("Real Property")
    _add_blank_table(
        fmt,
        ["Property Address", "Value", "Equity", "Transfer Method"],
        4,
    )


def _section_financial_accounts(fmt: DocxFormatter) -> None:
    fmt.h2("Financial Accounts")
    _add_blank_table(
        fmt,
        ["Institution", "Type", "Value", "Owner", "Beneficiary Designation"],
        6,
    )


def _section_vehicles(fmt: DocxFormatter) -> None:
    fmt.h2("Vehicles")
    _add_blank_table(
        fmt,
        ["Year/Make/Model/Description", "VIN", "Value", "Owner", "Transfer"],
        4,
    )


def _section_insurance(fmt: DocxFormatter) -> None:
    fmt.h2("Insurance Policies")
    _add_blank_table(
        fmt,
        ["Company", "Policy #", "Benefit", "Insured", "Beneficiary"],
        4,
    )


def _section_pensions(fmt: DocxFormatter) -> None:
    fmt.h2("Pensions / Retirement")
    _add_blank_table(
        fmt,
        ["Source", "Type", "Value", "Owner", "Survivor Provision"],
        4,
    )


def _section_valuables(fmt: DocxFormatter) -> None:
    fmt.h2("Valuables")
    _add_blank_table(
        fmt,
        ["Item Description", "Value", "Owner", "Specific Bequest?"],
        4,
    )


def _section_beneficiary_shares(fmt: DocxFormatter) -> None:
    fmt.h2("Beneficiary Shares")
    rows = _blank_rows(6, 4)
    rows.append(["TOTAL", "", "100%", ""])
    fmt.add_table(
        ["Beneficiary Name", "Relationship", "Share %", "Conditions"],
        rows,
    )
    _set_row_heights(fmt)


def _section_specific_bequests(fmt: DocxFormatter) -> None:
    fmt.h2("Specific Bequests")
    _add_blank_table(
        fmt,
        ["Item", "Recipient", "Special Instructions"],
        4,
    )


def _section_withdrawal_schedule(fmt: DocxFormatter) -> None:
    fmt.h2("Withdrawal Schedule")
    _add_blank_table(
        fmt,
        ["Step", "Timing", "Percentage"],
        4,
    )


# ---------------------------------------------------------------------------
# Elections (checkbox section)
# ---------------------------------------------------------------------------


def _section_elections(fmt: DocxFormatter, *, trust_type: str = "joint") -> None:
    fmt.h2("Trust Elections")
    fmt.body("Please check the box next to your preferred option for each category.")

    if trust_type == "individual":
        fmt.body("Initial Trustee:")
        fmt.indent(f"{_EMPTY_BOX}  Grantor as sole Initial Trustee")
    else:
        _checkbox_group(
            fmt,
            "Initial Trustee",
            [
                "Both Husband and Wife as Co-Trustees",
                "Husband only",
                "Wife only",
            ],
        )

        _checkbox_group(
            fmt,
            "Property Classification",
            [
                "Communal (all property treated as joint marital property)",
                "Separate (each spouse's property tracked separately)",
            ],
        )

    _checkbox_group(
        fmt,
        "Tangible Personal Property Distribution",
        [
            "Equally among children",
            "Equally among all beneficiaries",
        ],
    )

    _checkbox_group(
        fmt,
        "Division Method for Tangible Property",
        [
            "Trustee decides",
            "Lottery among beneficiaries",
            "Sell and divide proceeds",
        ],
    )

    _checkbox_group(
        fmt,
        "Distribution Standard",
        [
            "HEMS (Health, Education, Maintenance, and Support)",
            "Broad (trustee discretion)",
        ],
    )

    _checkbox_group(
        fmt,
        "If a Beneficiary Dies Before Full Distribution",
        [
            "Per stirpes (passes to that beneficiary's descendants)",
            "Per stirpes of the grantors",
            "Redistribute among remaining beneficiaries",
        ],
    )

    _checkbox_group(
        fmt,
        "Remote Contingent Beneficiary",
        [
            "Pass by intestacy (state law default)",
            "Donate to charity (specify below)",
        ],
    )
    fmt.indent("Charity name: " + "_" * 50)

    _checkbox_group(
        fmt,
        "Retirement Account Strategy",
        [
            "Payable-on-death to spouse, then children",
            "Payable to trust",
            "Mix of POD and trust",
        ],
    )

    _checkbox_group(
        fmt,
        "Life Insurance Strategy",
        [
            "Spouse as primary beneficiary, then children",
        ],
    )

    if trust_type != "individual":
        _checkbox_group(
            fmt,
            "Surviving Spouse Amendment Rights",
            [
                "Full amendment power",
                "Limited amendment power",
                "Irrevocable after first death",
            ],
        )

    _checkbox_group(
        fmt,
        "Power of Appointment",
        [
            "General power of appointment",
            "Limited power of appointment",
            "No power of appointment",
        ],
    )

    _checkbox_group(
        fmt,
        "No-Contest Clause",
        [
            "Yes — include a no-contest clause",
            "No",
        ],
    )

    _checkbox_group(
        fmt,
        "Spendthrift Provision",
        [
            "Yes — include spendthrift protection",
            "No",
        ],
    )

    _checkbox_group(
        fmt,
        "Probate Coordination",
        [
            "Yes — include pour-over will coordination",
            "No",
        ],
    )

    _checkbox_group(
        fmt,
        "Estate Tax Portability",
        [
            "Yes — elect portability",
            "No",
        ],
    )

    _checkbox_group(
        fmt,
        "Trustee Bond Required",
        [
            "Yes",
            "No — waive bond requirement",
        ],
    )

    _checkbox_group(
        fmt,
        "Dispute Resolution",
        [
            "Mediation, then binding arbitration",
            "Court litigation",
        ],
    )

    _checkbox_group(
        fmt,
        "Trustee Compensation",
        [
            "Reasonable compensation",
            "No compensation",
        ],
    )


# ---------------------------------------------------------------------------
# Freeform text sections
# ---------------------------------------------------------------------------


def _section_statement_of_intent(fmt: DocxFormatter) -> None:
    fmt.h2("Statement of Intent")
    fmt.body(
        "Please describe in your own words the primary purpose and goals of this trust:"
    )
    fmt.blank(6)


def _section_personal_message(fmt: DocxFormatter) -> None:
    fmt.h2("Personal Message to Beneficiaries")
    fmt.body(
        "If you wish to include a personal message to your beneficiaries "
        "in the trust document, please write it here:"
    )
    fmt.blank(6)


def _section_additional_notes(fmt: DocxFormatter) -> None:
    fmt.h2("Additional Notes")
    fmt.body(
        "Any other information, special circumstances, or instructions "
        "for the attorney:"
    )
    fmt.blank(6)


# ---------------------------------------------------------------------------
# Signature block
# ---------------------------------------------------------------------------


def _section_signature(fmt: DocxFormatter) -> None:
    fmt.blank(2)
    fmt.body("Date: " + "_" * 50)
    fmt.blank()
    fmt.body("Client Signature: " + "_" * 50)
    fmt.blank()
    fmt.body("Client Signature: " + "_" * 50)
