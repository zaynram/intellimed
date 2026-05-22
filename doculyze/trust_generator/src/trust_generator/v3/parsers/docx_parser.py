"""v3 docx intake-questionnaire parser.

Parses a `.docx` Trust Intake Questionnaire into a copy of the
seed-initialized `TrustData`. Honors the post-promotion contract from
`promote_seed` spec §6.2.3 (no re-invocation of `promote_seed`; joint
`trust_type` / `marital_status` mutation per `_resolve_captions`).

Public surface:
    parse_docx(filepath, seed_initialized) -> TrustData

This module is the docx leg of `parse_file`'s extension dispatch
(registered in `registry.py`). Two helpers extracted here are
cross-parser contract surfaces — the PDF parser in `pdf_parser.py`
imports them verbatim and the sequence.xml invariants pin their
signatures:

- ``_apply_post_promotion_protocol(result, parsed_trust_type, parsed_marital_status)``
  (cycle 4b, spec §5.3 step 4): post-promotion mutation +
  co_grantor materialization/dematerialization.
- ``_apply_post_merge_resolution(result, exclusions_string)``
  (cycle 6, spec §5.3 step 6): disinheritance resolution (§5.4.10)
  plus CorporateTrustee discrimination (§5.4.9).
"""

from __future__ import annotations

import logging
import re
from decimal import Decimal
from pathlib import Path
from typing import NamedTuple

from docx import Document  # type: ignore[import-untyped]

from trust_generator.v3.parsers.coercion import (
    _to_date,
    _to_decimal,
    _to_person_reference,
)
from trust_generator.v3.schema import (
    BeneficiaryShare,
    Child,
    CorporateTrustee,
    GrantorInfo,
    MaritalStatus,
    OtherBeneficiary,
    PersonReference,
    SuccessorTrustee,
    TrustData,
    TrustType,
    _resolve_captions,
    promote_seed,
)

# `promote_seed` is imported but NEVER called by this module. The import
# exists so that `unittest.mock.patch("trust_generator.v3.parsers.docx_parser.
# promote_seed")` resolves to a real attribute (spec §4 P1; pinned by
# test_parser_never_reinvokes_promote_seed). Do not remove without also
# removing the test.
_ = promote_seed  # type: ignore[unused-ignore]  # patch-site preservation

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Constants — checkbox / label / list-header / entity-suffix maps
# ---------------------------------------------------------------------------
#
# The extractor walks `doc.tables` row-by-row. Each row is dispatched to one
# of three handlers:
#
#   1. Checkbox row (col-0 check marker + col-1 known phrasing): emit a
#      typed enum value into the dedicated TrustType/MaritalStatus slots.
#   2. List header row (col-0 token matches a known header in
#      `_LIST_TABLE_HEADERS`): switch the table into a list-extraction
#      mode and consume subsequent rows as list entries.
#   3. Label/value row (col-0 in `_LABEL_KEY_MAP`): store col-1 text under
#      the mapped dotted key in `flat`.
#
# Disambiguation is by content, not position: a single `for table` pass
# inspects each row against each handler in priority order (header > checkbox
# > label) so the most-specific signal wins.

_CHECK_PREFIXES: tuple[str, ...] = (
    "X",
    "x",
    "☑",  # ☑
    "☒",  # ☒
    "[X]",
    "[x]",
)

_TRUST_TYPE_CHECKBOX_MAP: dict[str, TrustType] = {
    "this is a joint trust": TrustType.JOINT,
    "this is a individual trust": TrustType.INDIVIDUAL,
}

_MARITAL_STATUS_CHECKBOX_MAP: dict[str, MaritalStatus] = {
    "grantor is married": MaritalStatus.MARRIED,
    "grantor is unmarried": MaritalStatus.UNMARRIED,
}

_LABEL_KEY_MAP: dict[str, str] = {
    "Grantor full legal name": "grantor.full_legal_name",
    "Grantor date of birth": "grantor.date_of_birth",
    "Co-grantor full legal name": "co_grantor.full_legal_name",
}

# Header tokens (col-0 of the first row of a list-shaped table). Each maps
# to a section identifier consumed by the list-extraction dispatcher below.
_LIST_TABLE_HEADERS: dict[str, str] = {
    "Children": "children",
    "Successor Trustees": "successor_trustees",
    "Beneficiary Shares": "beneficiary_shares",
    "Other Beneficiaries": "other_beneficiaries",
}

# Paragraph prefix for the v2 exclusions text block (parser-internal carrier
# per F3 finding; not stored on `result`).
_EXCLUSIONS_PARAGRAPH_PREFIX = "Exclusions:"

# Spec §5.4.9 entity-name suffix heuristic. Conservative — known limitation:
# a natural person with surname "Bank" (e.g. "John Bank") is mis-typed. The
# INFO log inside `_apply_post_merge_resolution` is the operator-side
# recovery surface; the structural fix is the v3-questionnaire "entity?"
# checkbox (spec §9 Q4).
_ENTITY_NAME_PATTERN = re.compile(
    r"\b(Bank|Trust Company|Trust Department|N\.A\.|LLC|LLP|"
    r"Corporation|Corp\.|Inc\.|Insurance Co)\b",
    re.IGNORECASE,
)


def _is_check_marker(text: str) -> bool:
    """Return True if `text` looks like a checked-checkbox marker cell."""
    stripped = text.strip()
    return stripped in _CHECK_PREFIXES or any(
        stripped.startswith(p) for p in _CHECK_PREFIXES if len(p) > 1
    )


def _is_entity_name(name: str) -> bool:
    """§5.4.9 heuristic: does the name look like a corporate-trustee entity?

    Multi-token entity detection only; the one-token trap in
    ``coercion._to_person_reference`` handles bare entity-names like
    ``"AcmeCorp"`` separately. This helper is re-applied per trustee entry
    inside ``_apply_post_merge_resolution``. The boundary between this
    multi-token suffix heuristic (§5.4.9) and ``_to_person_reference``'s
    one-token trap (§5.4.4) is the layer split between the two helpers.
    """
    return bool(_ENTITY_NAME_PATTERN.search(name))


# ---------------------------------------------------------------------------
# Flat-key extraction
# ---------------------------------------------------------------------------


class _Extracted(NamedTuple):
    """Structured return of ``_extract_flat``.

    The shape grows monotonically across cycles 4b/5/6; using a NamedTuple
    keeps each slot strongly typed at consumer sites without value-narrowing
    gymnastics. ``_extract_flat`` is parser-internal — not a cross-plan
    contract — so the shape is implementation latitude.
    """

    flat: dict[str, str]
    exclusions: str
    parsed_trust_type: TrustType | None
    parsed_marital_status: MaritalStatus | None
    children: list[tuple[str, str]]
    successor_trustees: list[str]
    other_beneficiaries: list[str]
    beneficiary_shares: list[tuple[str, str]]


def _extract_flat(doc: object) -> _Extracted:
    """Walk ``doc.tables`` and ``doc.paragraphs`` into the parser's flat IR.

    Single-pass table walk dispatches each row to one of three handlers
    (header > checkbox > label) per the priority comment in the module-
    level constants section. Paragraphs are scanned for the v2 exclusions
    text-block prefix.
    """
    flat: dict[str, str] = {}
    exclusions = ""
    parsed_trust_type: TrustType | None = None
    parsed_marital_status: MaritalStatus | None = None
    children: list[tuple[str, str]] = []
    successor_trustees: list[str] = []
    other_beneficiaries: list[str] = []
    beneficiary_shares: list[tuple[str, str]] = []

    for table in doc.tables:  # type: ignore[attr-defined]
        rows = list(table.rows)
        if not rows:
            continue

        # Header dispatch: check the first row's col-0 against
        # _LIST_TABLE_HEADERS. On match, consume subsequent rows as list
        # entries and skip the rest of this table's other handlers.
        first_cells = rows[0].cells
        header_text = first_cells[0].text.strip() if first_cells else ""
        if header_text in _LIST_TABLE_HEADERS:
            section = _LIST_TABLE_HEADERS[header_text]
            for row in rows[1:]:
                cells = row.cells
                if not cells:
                    continue
                c0 = cells[0].text.strip()
                if not c0:
                    continue
                if section == "children":
                    dob = cells[1].text.strip() if len(cells) > 1 else ""
                    children.append((c0, dob))
                elif section == "successor_trustees":
                    successor_trustees.append(c0)
                elif section == "beneficiary_shares":
                    share = cells[1].text.strip() if len(cells) > 1 else ""
                    beneficiary_shares.append((c0, share))
                elif section == "other_beneficiaries":
                    other_beneficiaries.append(c0)
            continue

        # Otherwise: scan every row of this table for checkbox / label rows.
        matched = False
        for row in rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            c0 = cells[0].text.strip()
            c1 = cells[1].text.strip()

            # Checkbox row: column 0 is a check marker, column 1 is a v2.2
            # phrasing. First-match-wins; already-set fields are NOT
            # overwritten (deterministic earliest-match across multiple
            # checked options; surfaces a questionnaire authoring error
            # that the downstream diagnostic engine catches).
            if c0 and _is_check_marker(c0) and c1:
                matched = True
                c1_lower = c1.lower()
                if parsed_trust_type is None:
                    for phrase, tt in _TRUST_TYPE_CHECKBOX_MAP.items():
                        if phrase in c1_lower:
                            parsed_trust_type = tt
                            break
                if parsed_marital_status is None:
                    for phrase, ms in _MARITAL_STATUS_CHECKBOX_MAP.items():
                        if phrase in c1_lower:
                            parsed_marital_status = ms
                            break
                continue

            # Label/value row: column 0 is a known label, column 1 is the
            # value.
            if c0 in _LABEL_KEY_MAP:
                matched = True
                flat[_LABEL_KEY_MAP[c0]] = c1

        if len(rows) > 1 and not matched:
            log.warning(
                "table with %d rows matched no handler (header/checkbox/label); "
                "first-cell text: %r — possible template drift",
                len(rows),
                header_text[:80],
            )

    # Paragraph scan for the v2 exclusions text block. The prefix is
    # stripped before storing so the value is the bare token-list text.
    for para in doc.paragraphs:  # type: ignore[attr-defined]
        text = para.text.strip()
        if text.startswith(_EXCLUSIONS_PARAGRAPH_PREFIX):
            exclusions = text[len(_EXCLUSIONS_PARAGRAPH_PREFIX) :].strip()
            break

    return _Extracted(
        flat=flat,
        exclusions=exclusions,
        parsed_trust_type=parsed_trust_type,
        parsed_marital_status=parsed_marital_status,
        children=children,
        successor_trustees=successor_trustees,
        other_beneficiaries=other_beneficiaries,
        beneficiary_shares=beneficiary_shares,
    )


# ---------------------------------------------------------------------------
# Cross-parser helpers
# ---------------------------------------------------------------------------


def _apply_post_promotion_protocol(
    result: TrustData,
    parsed_trust_type: TrustType | None,
    parsed_marital_status: MaritalStatus | None,
) -> None:
    """Apply the §5.3 step 4 trust_type / marital_status mutation in place.

    **Signature contract (binding across parsers).** Consumed by both
    ``parse_docx`` in this module and ``parse_pdf`` in ``pdf_parser.py``
    (sibling-plan ``pdf`` cycle 7). The parameter names and order are
    pinned by sequence.xml's ``post-promotion-protocol-sig`` invariant:
    ``['result', 'parsed_trust_type', 'parsed_marital_status']``. Any
    change to this signature requires coordinated re-execution of both
    consuming cycles.

    **None-gate (F1, plan-review pass 2).** ``None`` for either parsed
    argument means "no mutation requested" — the seed-initialized value
    persists. The gate is load-bearing because ``trust_type`` is a
    required schema field and assigning ``None`` would breach Pydantic
    validation.

    **Ordering rule (spec §5.3 step 4).** Apply ``trust_type`` first
    (captions depend on it), then ``marital_status``; compute
    ``co_grantor`` materialization / dematerialization ONCE after both
    fields have settled.

    **Co_grantor protocol (spec §5.3 step 4 as amended by chore #37,
    2026-05-18 — lead-approved commit 2bc05da).**

    - Materialize: post-mutation state requires ``co_grantor`` AND
      ``result.co_grantor is None`` → ``result.co_grantor = GrantorInfo()``.
    - Dematerialize: post-mutation state requires no ``co_grantor`` AND
      ``result.co_grantor`` is field-equal to ``GrantorInfo()`` (no field
      populated beyond schema defaults; Pydantic v2 BaseModel field-equality
      makes this deterministic and is pinned by
      ``test_grantor_info_default_constructor_equality_is_deterministic``
      in cycle 5) → ``result.co_grantor = None``.
    - Preserve (implicit fallback): any populated ``co_grantor`` is left
      untouched per the bounded-context translation invariant; the data
      is meaningful and the parser must not drop it.
    """
    if parsed_trust_type is not None and parsed_trust_type != result.trust_id.trust_type:
        result.trust_id.trust_type = parsed_trust_type
        new_grantor_caption, new_co_grantor_caption = _resolve_captions(parsed_trust_type)
        result.trust_id.grantor_caption = new_grantor_caption
        result.trust_id.co_grantor_caption = new_co_grantor_caption

    if (
        parsed_marital_status is not None
        and parsed_marital_status != result.trust_id.marital_status
    ):
        result.trust_id.marital_status = parsed_marital_status

    should_have_co_grantor = (
        result.trust_id.trust_type == TrustType.JOINT
        or result.trust_id.marital_status == MaritalStatus.MARRIED
    )
    if should_have_co_grantor and result.co_grantor is None:
        result.co_grantor = GrantorInfo()
    elif not should_have_co_grantor and result.co_grantor == GrantorInfo():
        result.co_grantor = None


def _apply_post_merge_resolution(
    result: TrustData,
    exclusions_string: str,
) -> None:
    """Apply the §5.3 step 6 post-merge resolution passes.

    **Signature contract (binding across parsers).** Consumed by both
    ``parse_docx`` and ``parse_pdf`` (sibling-plan ``pdf`` cycle 7).
    Parameter names and order are pinned by sequence.xml's
    ``post-merge-resolution-sig`` invariant:
    ``['result', 'exclusions_string']`` (2-arg form). The 3-arg form some
    earlier plan iterations carried is OBSOLETE — peer-DM-aligned with the
    pdf teammate on 2026-05-18, the CorporateTrustee discrimination moved
    INSIDE this helper via ``_is_entity_name`` re-applied per trustee
    entry, replacing a caller-side flags carrier.

    **Pass 1 — Disinheritance resolution (§5.4.10).**

    1. Tokenize ``exclusions_string`` on commas, semicolons, and newlines.
    2. For each token, attempt case-insensitive substring match against
       beneficiaries in **fixed iteration order** (F2 finding):
       ``children`` → ``descendants`` → ``other_beneficiaries``. Within
       each list, Pydantic insertion order applies.
    3. First match wins. Secondary matches (later in the iteration order)
       are logged as a single WARNING naming both the chosen and
       secondary candidates (F4 finding) — they do NOT change the chosen
       target.
    4. On match: set the matched beneficiary's ``disinherit=True`` and
       ``disinherit_reason=token``.
    5. On no match: append ``PersonReference(full_legal_name=token)`` to
       ``result.external_exclusions`` and set
       ``result.external_exclusion_reasons[token] = token``.

    **Pass 2 — CorporateTrustee discrimination (§5.4.9).** Iterate
    ``result.successor_trustees``; for each entry, re-apply
    ``_is_entity_name`` to its ``full_legal_name``. On match, reconstruct
    the entry as ``CorporateTrustee(is_entity=True, full_legal_name="",
    entity_name=name)`` and emit an INFO log so the paralegal can review.

    **F3 note.** ``exclusions_string`` is function-local — NOT a field on
    ``result``. v3's ``TrustData`` has no ``text_blocks.exclusions`` field
    per ``modified_surfaces``; the extractor captures the string during
    step 2 and threads it as an argument.
    """
    # Pass 1: disinheritance resolution.
    if exclusions_string:
        tokens = [
            tok.strip()
            for tok in re.split(r"[,;\n]", exclusions_string)
            if tok.strip()
        ]
        # F2 (plan-review pass 2): fixed iteration order.
        iteration_buckets: list[tuple[str, list]] = [  # type: ignore[type-arg]
            ("children", result.children),
            ("descendants", result.descendants),
            ("other_beneficiaries", result.other_beneficiaries),
        ]
        for token in tokens:
            chosen: tuple[str, int, str] | None = None  # (bucket, idx, name)
            secondary_matches: list[tuple[str, str]] = []  # (bucket, name)
            token_lower = token.lower()
            for bucket_name, bucket in iteration_buckets:
                for idx, beneficiary in enumerate(bucket):
                    if token_lower in beneficiary.full_legal_name.lower():
                        if chosen is None:
                            chosen = (bucket_name, idx, beneficiary.full_legal_name)
                        else:
                            secondary_matches.append(
                                (bucket_name, beneficiary.full_legal_name)
                            )
            if chosen is None:
                # §5.4.10 step 4: unmatched token routes to
                # external_exclusions. PersonReference's two-token validator
                # accepts any 2+-token name; one-token tokens are an edge
                # case the spec does not cover and that the existing tests
                # do not exercise.
                result.external_exclusions.append(
                    PersonReference(full_legal_name=token)
                )
                result.external_exclusion_reasons[token] = token
            else:
                bucket_name, idx, chosen_name = chosen
                bucket_lookup = {name: bucket for name, bucket in iteration_buckets}
                bucket = bucket_lookup[bucket_name]
                bucket[idx].disinherit = True
                bucket[idx].disinherit_reason = token
                if secondary_matches:
                    secondary_names = ", ".join(name for _, name in secondary_matches)
                    log.warning(
                        "Disinheritance multi-match for token %r: chose %r "
                        "(iteration-order first; bucket=%s); secondary "
                        "candidates: %s",
                        token,
                        chosen_name,
                        bucket_name,
                        secondary_names,
                    )

    # Pass 2: CorporateTrustee discrimination (re-detect from preserved
    # name string; signature-coupling-free across docx / pdf consumers).
    if result.successor_trustees:
        new_trustees: list[SuccessorTrustee | CorporateTrustee] = []
        for trustee in result.successor_trustees:
            if not isinstance(trustee, CorporateTrustee) and _is_entity_name(
                trustee.full_legal_name
            ):
                new_trustees.append(
                    CorporateTrustee(
                        is_entity=True,
                        full_legal_name="",
                        entity_name=trustee.full_legal_name,
                    )
                )
                log.info(
                    "Discriminated %r as CorporateTrustee per §5.4.9 heuristic",
                    trustee.full_legal_name,
                )
            else:
                new_trustees.append(trustee)
        result.successor_trustees = new_trustees


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def parse_docx(filepath: Path, seed_initialized: TrustData) -> TrustData:
    """Parse a Trust Intake Questionnaire .docx INTO a copy of seed_initialized.

    The seed_initialized argument is required (no default) to make the
    post-promotion contract loud at every call site (spec §5.2). The
    return value is a deepcopied-then-filled TrustData; the caller's
    seed_initialized is field-level equal before and after this call
    (spec §4 P3, asserted by `test_parse_docx_smoke`).
    """
    if not filepath.exists():
        raise FileNotFoundError(filepath)

    result = seed_initialized.model_copy(deep=True)
    doc = Document(str(filepath))

    extracted = _extract_flat(doc)

    # Step 4 (post-promotion). The helper signature is binding across
    # parsers (pdf sibling imports it verbatim per sequence.xml's
    # post-promotion-protocol-sig invariant).
    _apply_post_promotion_protocol(
        result,
        extracted.parsed_trust_type,
        extracted.parsed_marital_status,
    )

    # Step 5 (remaining mutations) — field-by-field assignment with
    # coercion at each call site per spec §5.4.
    flat = extracted.flat

    if "grantor.full_legal_name" in flat:
        result.grantor.full_legal_name = flat["grantor.full_legal_name"]
    if "grantor.date_of_birth" in flat:
        result.grantor.date_of_birth = _to_date(flat["grantor.date_of_birth"])

    if "co_grantor.full_legal_name" in flat:
        # §5.4.4 placeholder-prefix stripping happens inside
        # _to_person_reference. The result's co_grantor was materialized
        # by promote_seed for (JT/MR) seeds OR by
        # _apply_post_promotion_protocol when the post-mutation state
        # required it; if it is still None (post-state requires no
        # co_grantor), drop the assignment.
        co_grantor_ref = _to_person_reference(flat["co_grantor.full_legal_name"])
        if result.co_grantor is not None:
            result.co_grantor.full_legal_name = co_grantor_ref.full_legal_name
            result.co_grantor.is_entity = co_grantor_ref.is_entity
            result.co_grantor.entity_name = co_grantor_ref.entity_name

    # Children: one Child per (name, dob) row. Date coercion via _to_date.
    for child_name, child_dob in extracted.children:
        ref = _to_person_reference(child_name)
        result.children.append(
            Child(
                full_legal_name=ref.full_legal_name,
                is_entity=ref.is_entity,
                entity_name=ref.entity_name,
                date_of_birth=_to_date(child_dob),
            )
        )

    # Successor trustees: construct as SuccessorTrustee uniformly. The
    # §5.4.9 entity discrimination + CorporateTrustee reconstruction
    # happens later inside _apply_post_merge_resolution (per the
    # 2026-05-18 docx ↔ pdf peer-DM contract — keeps the discrimination
    # in one place and the helper signature parser-agnostic).
    for trustee_name in extracted.successor_trustees:
        ref = _to_person_reference(trustee_name)
        result.successor_trustees.append(
            SuccessorTrustee(
                full_legal_name=ref.full_legal_name,
                is_entity=ref.is_entity,
                entity_name=ref.entity_name,
            )
        )

    # Other beneficiaries.
    for ob_name in extracted.other_beneficiaries:
        ref = _to_person_reference(ob_name)
        result.other_beneficiaries.append(
            OtherBeneficiary(
                full_legal_name=ref.full_legal_name,
                is_entity=ref.is_entity,
                entity_name=ref.entity_name,
            )
        )

    # Beneficiary shares: §5.4.2 share-percent branch drops rows whose
    # share-percent fails to parse to a non-zero Decimal (Decision log
    # #11). _to_decimal returns None on parse failure and Decimal(0) on
    # legitimate zero input; the parser layer drops on either — a 0%
    # share is as meaningless as an unparseable one. An emitted warning
    # naming the dropped row supplements the coercion helper's own
    # "could not parse decimal" warning.
    for share_name, share_pct_str in extracted.beneficiary_shares:
        share_pct = _to_decimal(share_pct_str)
        if share_pct is None or share_pct == Decimal(0):
            log.warning(
                "Dropping beneficiary_shares row for %r: could not parse "
                "share-percent %r per §5.4.2",
                share_name,
                share_pct_str,
            )
            continue
        ref = _to_person_reference(share_name)
        result.beneficiary_shares.append(
            BeneficiaryShare(
                recipient_external=ref,
                share_percent=share_pct,
            )
        )

    # Step 6 (post-merge resolution). The helper signature is binding
    # across parsers (pdf sibling imports it verbatim per sequence.xml's
    # post-merge-resolution-sig invariant).
    _apply_post_merge_resolution(result, extracted.exclusions)

    return result
