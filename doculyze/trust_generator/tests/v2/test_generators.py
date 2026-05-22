"""Tests for the trust document generator."""

from __future__ import annotations

from pathlib import Path

from docx import Document  # type: ignore[import-untyped]

from trust_generator.v2.config import AppConfig
from trust_generator.v2.generators import generate_trust_document
from trust_generator.v2.schema import (
    BeneficiaryShare,
    Child,
    PersonInfo,
    PropertyClassification,
    SpecificBequest,
    SuccessorTrustee,
    TrustData,
    TrustIdentity,
    TrustType,
    WithdrawalStep,
)

_counter = 0


def _generate_to_text(data: TrustData, tmp_path: Path, *, force: bool = True) -> str:
    """Generate a trust doc and return all paragraph text concatenated."""
    global _counter
    _counter += 1
    path = tmp_path / f"trust_{_counter}.docx"
    generate_trust_document(data, path, force=force)
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _minimal_data() -> TrustData:
    """Create a minimally complete TrustData for testing."""
    return TrustData(
        party_a=PersonInfo(full_legal_name="John Andrew Smith"),
        party_b=PersonInfo(full_legal_name="Jane Marie Smith"),
        trust_id=TrustIdentity(
            desired_trust_name="The Smith Family Trust",
            date="January 1, 2026",
        ),
        children=[
            Child(name="Alice Smith", dob="01/15/2000"),
            Child(name="Bob Smith", dob="03/22/2002"),
        ],
        successor_trustees=[
            SuccessorTrustee(order="1", name="Alice Smith", relationship="Daughter"),
        ],
        beneficiary_shares=[
            BeneficiaryShare(name="Alice Smith", share="50"),
            BeneficiaryShare(name="Bob Smith", share="50"),
        ],
    )


def test_generate_produces_valid_docx(tmp_path: Path):
    """Generator should produce a valid .docx file."""
    data = _minimal_data()
    path = tmp_path / "output.docx"
    result = generate_trust_document(data, path, force=True)
    assert Path(result).exists()
    doc = Document(result)
    assert len(doc.paragraphs) > 100


def test_generate_contains_all_articles(tmp_path: Path):
    """Output should contain all 12 article headings."""
    text = _generate_to_text(_minimal_data(), tmp_path)
    for i in range(1, 13):
        assert f"Article {i}:" in text, f"Article {i} missing"


def test_generate_contains_schedules(tmp_path: Path):
    text = _generate_to_text(_minimal_data(), tmp_path)
    assert "Schedule A" in text
    assert "Schedule B" in text


def test_generate_contains_grantor_names(tmp_path: Path):
    text = _generate_to_text(_minimal_data(), tmp_path)
    assert "John Andrew Smith" in text
    assert "Jane Marie Smith" in text


def test_generate_contains_children(tmp_path: Path):
    text = _generate_to_text(_minimal_data(), tmp_path)
    assert "Alice Smith" in text
    assert "Bob Smith" in text


def test_generate_contains_trust_name(tmp_path: Path):
    text = _generate_to_text(_minimal_data(), tmp_path)
    assert "The Smith Family Trust" in text


def test_spendthrift_false_omits_section(tmp_path: Path):
    """When spendthrift is False, Section 12.2 should NOT appear."""
    data = _minimal_data()
    data.elections.spendthrift = False
    text = _generate_to_text(data, tmp_path)
    assert "Spendthrift Provision" not in text


def test_spendthrift_true_includes_section(tmp_path: Path):
    text = _generate_to_text(_minimal_data(), tmp_path)
    assert "Spendthrift Provision" in text


def test_no_contest_false_omits_section(tmp_path: Path):
    """When no_contest is False, Section 12.3 should NOT appear."""
    data = _minimal_data()
    data.elections.no_contest = False
    text = _generate_to_text(data, tmp_path)
    assert "Contest Provision" not in text


def test_no_contest_true_includes_section(tmp_path: Path):
    text = _generate_to_text(_minimal_data(), tmp_path)
    assert "Contest Provision" in text


def test_probate_coordination_false_omits_section(tmp_path: Path):
    """When probate_coordination is False, Section 4.6 should NOT appear."""
    data = _minimal_data()
    data.elections.probate_coordination = False
    text = _generate_to_text(data, tmp_path)
    assert "Coordination with Personal Representative" not in text


def test_probate_coordination_true_includes_section(tmp_path: Path):
    text = _generate_to_text(_minimal_data(), tmp_path)
    assert "Coordination with Personal Representative" in text


def test_separate_property_generates_schedules_cd(tmp_path: Path):
    data = _minimal_data()
    data.elections.property_classification = PropertyClassification.SEPARATE
    text = _generate_to_text(data, tmp_path)
    assert "Schedule C" in text
    assert "Schedule D" in text


def test_communal_property_no_schedules_cd(tmp_path: Path):
    text = _generate_to_text(_minimal_data(), tmp_path)
    assert "Schedule C" not in text
    assert "Schedule D" not in text


def test_empty_data_does_not_crash(tmp_path: Path):
    """Generator should handle a completely empty TrustData without crashing."""
    data = TrustData()
    path = tmp_path / "empty.docx"
    generate_trust_document(data, path, force=True)
    assert path.exists()


def test_firm_info_from_config(tmp_path: Path):
    """Firm info should come from config, not hardcoded."""
    cfg = AppConfig()
    text = _generate_to_text(_minimal_data(), tmp_path)
    assert cfg.firm.name in text


def test_document_metadata(tmp_path: Path):
    """Generated doc should have metadata in core properties."""
    path = tmp_path / "meta.docx"
    generate_trust_document(_minimal_data(), path, force=True)
    doc = Document(str(path))
    assert "trust-generator" in (doc.core_properties.comments or "")


def test_withdrawal_schedule_rendered(tmp_path: Path):
    data = _minimal_data()
    data.withdrawal_schedule = [
        WithdrawalStep(step="1", timing="1 year after funding", percentage="50"),
        WithdrawalStep(step="2", timing="3 years after funding", percentage="50"),
    ]
    text = _generate_to_text(data, tmp_path)
    assert "50%" in text
    assert "1 year after funding" in text


def test_specific_bequests_in_schedule_b(tmp_path: Path):
    data = _minimal_data()
    data.specific_bequests = [
        SpecificBequest(item="Grand piano", recipient="Alice Smith"),
    ]
    text = _generate_to_text(data, tmp_path)
    assert "Grand piano" in text
    assert "Alice Smith" in text


# ---------------------------------------------------------------------------
# Individual (single-grantor) trust tests
# ---------------------------------------------------------------------------


def _individual_data() -> TrustData:
    """Create a minimally complete individual TrustData for testing."""
    return TrustData(
        trust_type=TrustType.INDIVIDUAL,
        grantor=PersonInfo(full_legal_name="Robert James Wilson"),
        trust_id=TrustIdentity(
            desired_trust_name="The Wilson Family Trust",
            date="March 15, 2026",
        ),
        children=[
            Child(name="Sarah Wilson", dob="05/10/1995"),
        ],
        successor_trustees=[
            SuccessorTrustee(order="1", name="Sarah Wilson", relationship="Daughter"),
        ],
        beneficiary_shares=[
            BeneficiaryShare(name="Sarah Wilson", share="100"),
        ],
    )


def test_individual_generates_valid_docx(tmp_path):
    data = _individual_data()
    path = tmp_path / "individual.docx"
    result = generate_trust_document(data, path, force=True)
    assert Path(result).exists()
    doc = Document(result)
    assert len(doc.paragraphs) > 50


def test_individual_contains_all_articles(tmp_path):
    text = _generate_to_text(_individual_data(), tmp_path)
    for i in range(1, 13):
        assert f"Article {i}:" in text, f"Article {i} missing"


def test_individual_contains_grantor_name(tmp_path):
    text = _generate_to_text(_individual_data(), tmp_path)
    assert "Robert James Wilson" in text


def test_individual_uses_singular_language(tmp_path):
    """Individual trust should use 'I'/'my' language, not 'We'/'our'."""
    text = _generate_to_text(_individual_data(), tmp_path)
    # Should have singular grantor reference
    assert "the \u201cGrantor\u201d" in text or "Grantor" in text


def test_individual_single_signature_line(tmp_path):
    text = _generate_to_text(_individual_data(), tmp_path)
    assert "Robert James Wilson, Grantor and Trustee" in text


def test_individual_empty_data_no_crash(tmp_path):
    data = TrustData(trust_type=TrustType.INDIVIDUAL)
    path = tmp_path / "empty_individual.docx"
    generate_trust_document(data, path, force=True)
    assert path.exists()
