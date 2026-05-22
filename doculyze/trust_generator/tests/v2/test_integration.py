"""Integration tests and edge-case tests for the trust-generator pipeline."""

from __future__ import annotations

from functools import partial
from pathlib import Path

import pytest
from docx import Document  # type: ignore[import-untyped]

from trust_generator.v2.config import load_config
from trust_generator.v2.generators import generate_trust_document
from trust_generator.v2.parsers import parse_docx, parse_json
from trust_generator.v2.schema import (
    BeneficiaryShare,
    Child,
    DistributionStandard,
    Elections,
    FinancialAccount,
    InitialTrustee,
    InsurancePolicy,
    Pension,
    PersonInfo,
    PowerOfAppointment,
    PropertyClassification,
    RealProperty,
    RemoteContingent,
    SuccessorTrustee,
    SurvivingAmendment,
    TangibleDistribution,
    TrustData,
    TrustIdentity,
    TrustType,
    Valuable,
    Vehicle,
    WithdrawalStep,
)
from trust_generator.v2.validators import Severity
from trust_generator.v2.validators import validate as _validate

validate = partial(_validate, config=load_config())
# ---------------------------------------------------------------------------
# Fixtures and helpers
# ---------------------------------------------------------------------------

ASSETS_DIR = Path(__file__).resolve().parents[2] / "assets"
QUESTIONNAIRE_PATH = ASSETS_DIR / "Trust_Intake_Questionnaire.docx"


def _generate_to_text(data: TrustData, tmp_path: Path, *, force: bool = True) -> str:
    """Generate a trust doc into tmp_path and return all paragraph text."""
    out = tmp_path / "output.docx"
    generate_trust_document(data, out, force=force)
    doc = Document(str(out))
    return "\n".join(p.text for p in doc.paragraphs)


def _minimal_data() -> TrustData:
    """TrustData with only grantor names filled in."""
    return TrustData(
        party_a=PersonInfo(full_legal_name="John Andrew Smith"),
        party_b=PersonInfo(full_legal_name="Jane Marie Smith"),
    )


def _complete_data() -> TrustData:
    """A reasonably complete TrustData for pipeline tests."""
    return TrustData(
        party_a=PersonInfo(full_legal_name="John Andrew Smith", ssn="123-45-6789"),
        party_b=PersonInfo(full_legal_name="Jane Marie Smith"),
        trust_id=TrustIdentity(
            desired_trust_name="The Smith Family Trust",
            date="January 1, 2026",
        ),
        children=[
            Child(name="Alice Smith", dob="01/15/2000", relationship="Daughter"),
            Child(name="Bob Smith", dob="03/22/2002", relationship="Son"),
        ],
        successor_trustees=[
            SuccessorTrustee(order="1", name="Alice Smith", relationship="Daughter"),
        ],
        beneficiary_shares=[
            BeneficiaryShare(name="Alice Smith", share="50"),
            BeneficiaryShare(name="Bob Smith", share="50"),
        ],
        real_property=[RealProperty(address="123 Main St", equity="$200,000")],
    )


# ===========================================================================
# End-to-end pipeline tests
# ===========================================================================


@pytest.mark.skipif(
    not QUESTIONNAIRE_PATH.exists(),
    reason="Trust_Intake_Questionnaire.docx not found in assets/",
)
def test_docx_to_trust_document(tmp_path: Path) -> None:
    """Parse the real questionnaire .docx, validate, generate, and verify output."""
    data = parse_docx(QUESTIONNAIRE_PATH)
    report = validate(data)
    # Even the blank questionnaire should be force-generatable (errors allowed)
    out = tmp_path / "trust_output.docx"
    generate_trust_document(data, out, force=True)

    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    for i in range(1, 13):
        assert f"Article {i}:" in text, f"Article {i} missing from generated output"
    # Validation report should exist and have findings
    assert len(report.findings) > 0


def test_json_round_trip_pipeline(tmp_path: Path) -> None:
    """Create TrustData -> dump to JSON -> parse JSON -> validate -> generate."""
    original = _complete_data()

    # Dump to JSON
    json_file = tmp_path / "intake.json"
    json_file.write_text(original.model_dump_json(indent=2), encoding="utf-8")

    # Parse back
    parsed = parse_json(json_file)
    assert parsed == original

    # Validate
    report = validate(parsed)
    assert report.can_generate is True

    # Generate
    out = tmp_path / "trust_output.docx"
    generate_trust_document(parsed, out, force=True)

    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    for i in range(1, 13):
        assert f"Article {i}:" in text, f"Article {i} missing"
    assert "John Andrew Smith" in text
    assert "Jane Marie Smith" in text


def test_parse_validate_generate_minimal(tmp_path: Path) -> None:
    """Minimal TrustData (just names) -> JSON -> parse -> validate -> generate with force."""
    data = _minimal_data()

    json_file = tmp_path / "minimal.json"
    json_file.write_text(data.model_dump_json(indent=2), encoding="utf-8")

    parsed = parse_json(json_file)
    report = validate(parsed)

    # Should have warnings but no errors (names are provided)
    assert report.can_generate is True
    assert len(report.warnings) > 0

    out = tmp_path / "minimal_trust.docx"
    generate_trust_document(parsed, out, force=True)
    assert out.exists()
    assert out.stat().st_size > 0


# ===========================================================================
# Parser edge cases
# ===========================================================================


@pytest.mark.skipif(
    not QUESTIONNAIRE_PATH.exists(),
    reason="Trust_Intake_Questionnaire.docx not found in assets/",
)
def test_docx_parser_returns_trust_data_type() -> None:
    """parse_docx must return exactly a TrustData instance, not a dict."""
    result = parse_docx(QUESTIONNAIRE_PATH)
    assert type(result) is TrustData


def test_json_parser_preserves_elections(tmp_path: Path) -> None:
    """Non-default elections should survive a JSON round-trip."""
    data = TrustData(
        party_a=PersonInfo(full_legal_name="Test Husband"),
        party_b=PersonInfo(full_legal_name="Test Wife"),
        elections=Elections(
            spendthrift=False,
            distribution_standard=DistributionStandard.BROAD,
            initial_trustee=InitialTrustee.PARTY_B,
            no_contest=False,
            probate_coordination=False,
            property_classification=PropertyClassification.SEPARATE,
            surviving_amendment=SurvivingAmendment.IRREVOCABLE,
            power_of_appointment=PowerOfAppointment.NONE,
            remote_contingent=RemoteContingent.CHARITY,
            remote_contingent_charity="Red Cross",
        ),
    )

    json_file = tmp_path / "elections.json"
    json_file.write_text(data.model_dump_json(indent=2), encoding="utf-8")

    restored = parse_json(json_file)

    assert restored.elections.spendthrift is False
    assert restored.elections.distribution_standard == DistributionStandard.BROAD
    assert restored.elections.initial_trustee == InitialTrustee.PARTY_B
    assert restored.elections.no_contest is False
    assert restored.elections.probate_coordination is False
    assert restored.elections.property_classification == PropertyClassification.SEPARATE
    assert restored.elections.surviving_amendment == SurvivingAmendment.IRREVOCABLE
    assert restored.elections.power_of_appointment == PowerOfAppointment.NONE
    assert restored.elections.remote_contingent == RemoteContingent.CHARITY
    assert restored.elections.remote_contingent_charity == "Red Cross"


def test_json_parser_preserves_all_asset_types(tmp_path: Path) -> None:
    """One item in each of the 6 asset categories should survive JSON round-trip."""
    data = TrustData(
        party_a=PersonInfo(full_legal_name="Asset Test"),
        real_property=[RealProperty(address="123 Main St", equity="$200,000")],
        financial_accounts=[
            FinancialAccount(institution="Chase", type="Checking", value="$50,000")
        ],
        vehicles=[Vehicle(description="2020 Toyota Camry", value="$25,000")],
        insurance_policies=[InsurancePolicy(company="State Farm", benefit="$500,000")],
        pensions=[Pension(source="Employer Corp", type="401k", value="$100,000")],
        valuables=[Valuable(description="Antique Clock", value="$5,000")],
    )

    json_file = tmp_path / "assets.json"
    json_file.write_text(data.model_dump_json(indent=2), encoding="utf-8")

    restored = parse_json(json_file)

    assert len(restored.real_property) == 1
    assert restored.real_property[0].address == "123 Main St"
    assert len(restored.financial_accounts) == 1
    assert restored.financial_accounts[0].institution == "Chase"
    assert len(restored.vehicles) == 1
    assert restored.vehicles[0].description == "2020 Toyota Camry"
    assert len(restored.insurance_policies) == 1
    assert restored.insurance_policies[0].company == "State Farm"
    assert len(restored.pensions) == 1
    assert restored.pensions[0].source == "Employer Corp"
    assert len(restored.valuables) == 1
    assert restored.valuables[0].description == "Antique Clock"


# ===========================================================================
# Validator edge cases
# ===========================================================================


def test_validator_beneficiary_shares_non_numeric() -> None:
    """Shares with non-numeric text should produce a warning about unparseable shares."""
    data = _complete_data()
    data.beneficiary_shares = [
        BeneficiaryShare(name="Alice Smith", share="fifty"),
        BeneficiaryShare(name="Bob Smith", share="fifty"),
    ]
    report = validate(data)
    share_warnings = [
        f
        for f in report.warnings
        if f.field_path == "beneficiary_shares" and "parsed" in f.message.lower()
    ]
    assert len(share_warnings) == 1


def test_validator_all_defaults_has_warnings() -> None:
    """TrustData with only grantor names should produce INFO findings for defaulted fields."""
    data = _minimal_data()
    report = validate(data)
    info_findings = [f for f in report.findings if f.severity == Severity.INFO]
    assert len(info_findings) > 0


def test_validator_charity_with_name_no_error() -> None:
    """CHARITY remote contingent WITH a charity name should NOT produce an error."""
    data = _complete_data()
    data.elections = Elections(
        remote_contingent=RemoteContingent.CHARITY,
        remote_contingent_charity="Red Cross Foundation",
    )
    report = validate(data)
    charity_errors = [
        f
        for f in report.errors
        if f.field_path == "elections.remote_contingent_charity"
    ]
    assert len(charity_errors) == 0


# ===========================================================================
# Generator edge cases
# ===========================================================================


def test_generator_all_elections_non_default(tmp_path: Path) -> None:
    """Generate with every election set to non-default values; verify no crash."""
    data = TrustData(
        party_a=PersonInfo(full_legal_name="John Andrew Smith"),
        party_b=PersonInfo(full_legal_name="Jane Marie Smith"),
        trust_id=TrustIdentity(
            desired_trust_name="The Smith Family Trust",
            date="January 1, 2026",
        ),
        children=[
            Child(name="Alice Smith", dob="01/15/2000", relationship="Daughter"),
        ],
        successor_trustees=[
            SuccessorTrustee(order="1", name="Alice Smith", relationship="Daughter"),
        ],
        beneficiary_shares=[
            BeneficiaryShare(name="Alice Smith", share="100"),
        ],
        elections=Elections(
            initial_trustee=InitialTrustee.PARTY_B,
            property_classification=PropertyClassification.SEPARATE,
            tangible_distribution=TangibleDistribution.EQUAL_BENEFICIARIES,
            distribution_standard=DistributionStandard.BROAD,
            surviving_amendment=SurvivingAmendment.IRREVOCABLE,
            power_of_appointment=PowerOfAppointment.NONE,
            remote_contingent=RemoteContingent.CHARITY,
            remote_contingent_charity="Local Food Bank",
            spendthrift=False,
            no_contest=False,
            probate_coordination=False,
            portability=False,
            trustee_bond=True,
        ),
    )
    text = _generate_to_text(data, tmp_path)

    # Non-default election effects
    assert "Jane Marie Smith" in text  # WIFE-only trustee name should appear
    assert "Schedule C" in text  # SEPARATE property
    assert "Schedule D" in text
    assert "any purpose" in text  # BROAD distribution
    assert "irrevocable" in text.lower()  # IRREVOCABLE surviving amendment
    assert "Local Food Bank" in text  # CHARITY contingent
    # Opted-out sections should NOT appear
    assert "Spendthrift Provision" not in text
    assert "Contest Provision" not in text
    assert "Coordination with Personal Representative" not in text


def test_generator_many_children(tmp_path: Path) -> None:
    """Generate with 10 children; verify all names appear in output."""
    children = [
        Child(name=f"Child_{i} Smith", dob=f"01/{i:02d}/2000", relationship="Child")
        for i in range(1, 11)
    ]
    data = TrustData(
        party_a=PersonInfo(full_legal_name="John Andrew Smith"),
        party_b=PersonInfo(full_legal_name="Jane Marie Smith"),
        children=children,
    )
    text = _generate_to_text(data, tmp_path)

    for i in range(1, 11):
        assert f"Child_{i} Smith" in text, f"Child_{i} Smith not found in output"
    assert "10 children" in text


def test_generator_many_assets(tmp_path: Path) -> None:
    """Generate with items in all 6 asset categories; verify all appear."""
    data = TrustData(
        party_a=PersonInfo(full_legal_name="John Andrew Smith"),
        party_b=PersonInfo(full_legal_name="Jane Marie Smith"),
        real_property=[RealProperty(address="123 Main St", equity="$200,000")],
        financial_accounts=[
            FinancialAccount(institution="Chase", type="Checking", value="$50,000")
        ],
        vehicles=[Vehicle(description="2020 Toyota Camry", value="$25,000")],
        insurance_policies=[InsurancePolicy(company="State Farm", benefit="$500,000")],
        pensions=[Pension(source="Employer Corp", type="401k", value="$100,000")],
        valuables=[Valuable(description="Antique Clock", value="$5,000")],
    )
    text = _generate_to_text(data, tmp_path)

    assert "123 Main St" in text
    assert "Chase" in text
    assert "Toyota Camry" in text
    assert "State Farm" in text
    assert "Employer Corp" in text
    assert "Antique Clock" in text


def test_generator_withdrawal_schedule_empty_vs_populated(tmp_path: Path) -> None:
    """Empty schedule should have placeholder; populated should have actual percentages."""
    # Empty schedule
    data_empty = TrustData(
        party_a=PersonInfo(full_legal_name="John Andrew Smith"),
        party_b=PersonInfo(full_legal_name="Jane Marie Smith"),
    )
    out_empty = tmp_path / "empty_schedule.docx"
    generate_trust_document(data_empty, out_empty, force=True)
    doc_empty = Document(str(out_empty))
    text_empty = "\n".join(p.text for p in doc_empty.paragraphs)

    assert "[SPECIFY SCHEDULE]" in text_empty

    # Populated schedule
    data_pop = TrustData(
        party_a=PersonInfo(full_legal_name="John Andrew Smith"),
        party_b=PersonInfo(full_legal_name="Jane Marie Smith"),
        withdrawal_schedule=[
            WithdrawalStep(step="1", timing="1 year after funding", percentage="33"),
            WithdrawalStep(step="2", timing="3 years after funding", percentage="33"),
            WithdrawalStep(step="3", timing="5 years after funding", percentage="34"),
        ],
    )
    out_pop = tmp_path / "populated_schedule.docx"
    generate_trust_document(data_pop, out_pop, force=True)
    doc_pop = Document(str(out_pop))
    text_pop = "\n".join(p.text for p in doc_pop.paragraphs)

    assert "[SPECIFY SCHEDULE]" not in text_pop
    assert "33%" in text_pop
    assert "34%" in text_pop
    assert "1 year after funding" in text_pop
    assert "3 years after funding" in text_pop
    assert "5 years after funding" in text_pop


# ===========================================================================
# Individual trust pipeline tests
# ===========================================================================


def test_individual_trust_json_pipeline(tmp_path: Path) -> None:
    """Individual trust: create -> JSON -> parse -> validate -> generate."""
    original = TrustData(
        trust_type=TrustType.INDIVIDUAL,
        grantor=PersonInfo(full_legal_name="Robert James Wilson", ssn="987-65-4321"),
        trust_id=TrustIdentity(
            desired_trust_name="The Wilson Family Trust",
            date="March 15, 2026",
        ),
        children=[
            Child(name="Sarah Wilson", dob="05/10/1995", relationship="Daughter"),
        ],
        successor_trustees=[
            SuccessorTrustee(order="1", name="Sarah Wilson", relationship="Daughter"),
        ],
        beneficiary_shares=[
            BeneficiaryShare(name="Sarah Wilson", share="100"),
        ],
        real_property=[RealProperty(address="456 Oak Ave", equity="$300,000")],
    )

    # Dump to JSON
    json_file = tmp_path / "individual_intake.json"
    json_file.write_text(original.model_dump_json(indent=2), encoding="utf-8")

    # Parse back
    parsed = parse_json(json_file)
    assert parsed.trust_type == TrustType.INDIVIDUAL
    assert parsed.grantor.full_legal_name == "Robert James Wilson"

    # Validate
    report = validate(parsed)
    assert report.can_generate is True

    # Generate
    out = tmp_path / "wilson_trust.docx"
    generate_trust_document(parsed, out, force=True)
    assert out.exists()

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    for i in range(1, 13):
        assert f"Article {i}:" in text, f"Article {i} missing"
    assert "Robert James Wilson" in text
    assert "The Wilson Family Trust" in text


# ===========================================================================
# PDF pipeline integration test
# ===========================================================================


try:
    import pypdf
    import reportlab

    _HAS_PDF_DEPS = True
except ImportError:
    _HAS_PDF_DEPS = False


@pytest.mark.skipif(not _HAS_PDF_DEPS, reason="pypdf/reportlab not installed")
def test_pdf_generate_fill_parse_validate_generate(tmp_path: Path) -> None:
    """Full PDF pipeline: generate blank PDF → fill → parse → validate → generate trust doc."""
    from pypdf import PdfReader, PdfWriter

    from trust_generator.v2.generators import generate_fillable_pdf
    from trust_generator.v2.parsers import parse_pdf

    # Step 1: Generate blank fillable PDF
    blank_pdf = tmp_path / "blank.pdf"
    generate_fillable_pdf(blank_pdf)
    assert blank_pdf.exists()

    # Step 2: Fill in fields programmatically
    reader = PdfReader(str(blank_pdf))
    writer = PdfWriter()
    writer.append(reader)

    fill_data = {
        "party_a.full_legal_name": "John Andrew Smith",
        "party_b.full_legal_name": "Jane Marie Smith",
        "trust_id.desired_trust_name": "The Smith Family Trust",
        "trust_id.date": "January 1, 2026",
        "trust_id.state_of_governing_law": "Illinois",
        "trust_id.county_of_execution": "Winnebago",
    }

    for page in writer.pages:
        writer.update_page_form_field_values(page, fill_data)

    filled_pdf = tmp_path / "filled.pdf"
    with open(filled_pdf, "wb") as f:
        writer.write(f)

    # Step 3: Parse
    data = parse_pdf(filled_pdf)
    assert data.party_a.full_legal_name == "John Andrew Smith"
    assert data.party_b.full_legal_name == "Jane Marie Smith"

    # Step 4: Validate
    report = validate(data)
    assert report.can_generate is True

    # Step 5: Generate trust document
    out = tmp_path / "smith_trust.docx"
    generate_trust_document(data, out, force=True)
    assert out.exists()

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "John Andrew Smith" in text
    assert "The Smith Family Trust" in text
    for i in range(1, 13):
        assert f"Article {i}:" in text
