"""Tests for the printable questionnaire generator."""

from __future__ import annotations

from pathlib import Path

from docx import Document  # type: ignore[import-untyped]

from trust_generator.v2.config import AppConfig, FirmConfig
from trust_generator.v2.generators import generate_printable_questionnaire


def _generate_and_read(
    tmp_path: Path, config: AppConfig | None = None
) -> tuple[str, str]:
    """Generate the questionnaire and return (path, full_text)."""
    path = tmp_path / "questionnaire.docx"
    generate_printable_questionnaire(str(path), config=config)
    doc = Document(str(path))

    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    full_text = "\n".join(parts)
    return str(path), full_text


def test_generates_valid_docx(tmp_path: Path) -> None:
    """Output file exists and is a valid .docx that python-docx can open."""
    path = tmp_path / "output.docx"
    result = generate_printable_questionnaire(str(path))
    assert Path(result).exists()
    doc = Document(result)
    assert len(doc.paragraphs) > 20


def test_contains_all_sections(tmp_path: Path) -> None:
    """Key section headings appear in the output."""
    _, text = _generate_and_read(tmp_path)
    expected_sections = [
        "Office Use",
        "Husband Information",
        "Wife Information",
        "Marriage Information",
        "Trust Information",
        "Children",
        "Successor Trustees",
        "Real Property",
        "Financial Accounts",
        "Vehicles",
        "Insurance Policies",
        "Pensions / Retirement",
        "Valuables",
        "Beneficiary Shares",
        "Specific Bequests",
        "Withdrawal Schedule",
        "Trust Elections",
        "Statement of Intent",
        "Personal Message to Beneficiaries",
        "Additional Notes",
    ]
    for section in expected_sections:
        assert section in text, f"Missing section: {section}"


def test_no_placeholder_text(tmp_path: Path) -> None:
    """Common placeholder strings must NOT appear anywhere in the output."""
    _, text = _generate_and_read(tmp_path)
    forbidden = [
        "e.g.,",
        "e.g. ",
        "MM/DD/YYYY",
        "XXX-XX-XXXX",
        "John Andrew Doe",
        "(enter ",
        "placeholder",
    ]
    for hint in forbidden:
        assert hint.lower() not in text.lower(), f"Placeholder text found: {hint!r}"


def test_contains_firm_name(tmp_path: Path) -> None:
    """Firm name from config appears in the output."""
    cfg = AppConfig(firm=FirmConfig(name="Test Law Firm LLP"))
    _, text = _generate_and_read(tmp_path, config=cfg)
    assert "Test Law Firm LLP" in text


def test_contains_checkbox_symbols(tmp_path: Path) -> None:
    """The empty checkbox symbol appears for election sections."""
    _, text = _generate_and_read(tmp_path)
    assert "\u2610" in text, "No empty checkbox symbol (☐) found"
    assert text.count("\u2610") >= 20


def test_individual_printable_questionnaire(tmp_path: Path) -> None:
    path = tmp_path / "individual_questionnaire.docx"
    result = generate_printable_questionnaire(path, trust_type="individual")
    assert Path(result).exists()
    doc = Document(result)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Grantor Information" in text
    assert "Husband Information" not in text
    assert "Wife Information" not in text
    assert "Marriage Information" not in text


def test_joint_printable_questionnaire_unchanged(tmp_path: Path) -> None:
    path = tmp_path / "joint_questionnaire.docx"
    result = generate_printable_questionnaire(path)
    doc = Document(result)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Husband Information" in text
    assert "Wife Information" in text


def test_individual_printable_hides_spousal_elections(tmp_path: Path) -> None:
    """Individual trust questionnaire hides joint-only election options."""
    path = tmp_path / "individual_elections.docx"
    result = generate_printable_questionnaire(path, trust_type="individual")
    doc = Document(result)
    text = "\n".join(p.text for p in doc.paragraphs)
    # Should show individual trustee option
    assert "Grantor as sole Initial Trustee" in text
    # Should NOT show joint-only options
    assert "Both Husband and Wife as Co-Trustees" not in text
    assert "Surviving Spouse Amendment Rights" not in text
    assert "Communal" not in text
    assert "Separate (each spouse" not in text


def test_joint_printable_keeps_all_election_options(tmp_path: Path) -> None:
    """Joint trust questionnaire keeps all election options."""
    path = tmp_path / "joint_elections.docx"
    result = generate_printable_questionnaire(path, trust_type="joint")
    doc = Document(result)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Both Husband and Wife as Co-Trustees" in text
    assert "Husband only" in text
    assert "Wife only" in text
    assert "Surviving Spouse Amendment Rights" in text
    assert "Communal" in text


def test_election_checkbox_text_is_canonical(tmp_path: Path) -> None:
    """Election checkbox text uses canonical labels for parser round-trip."""
    path = tmp_path / "canonical_elections.docx"
    result = generate_printable_questionnaire(
        path,
        trust_type="joint",
        party_a_label="Grantor A",
        party_b_label="Grantor B",
    )
    doc = Document(result)
    text = "\n".join(p.text for p in doc.paragraphs)
    # Election checkboxes should use canonical "Husband"/"Wife", not dynamic labels
    assert "Husband only" in text
    assert "Wife only" in text
    assert "Grantor A only" not in text
    assert "Grantor B only" not in text
