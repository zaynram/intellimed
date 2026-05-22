"""Programmatic .docx fixture builder for parser tests.

Decision log #20 (spec §6.7.1): synthetic docx fixtures are constructed
in-test via python-docx's ``Document()`` API. Checked-in fixture binaries
are explicitly NOT used because:

1. Programmatic construction couples fixture content to parser
   expectations directly in the test source — changes to parser
   table-detection logic can be reflected in the fixture builder in
   the same commit.
2. Binary fixtures drift silently when the parser's expectations change;
   the diff becomes opaque.

The builder emits ``.docx`` files into a caller-supplied ``tmp_path`` and
returns the path. Each kwarg controls one section of the v2.2
questionnaire layout; absent kwargs produce no content for that section.

The exact label strings and header phrasings emitted here MUST match the
labels, checkbox phrasings, and header tokens recognized by
``docx_parser._extract_flat``. Any drift between the two layers is a
fixture/parser-label coupling bug surfaced by the cycle-5/6 tests;
resolve by adjusting the fixture's emitted phrasing to match the
parser's expected label (the parser's labels are authoritative because
they come from the v2.2 questionnaire content).

**Header convention (cycle 6).** List-shaped tables (children,
successor_trustees, beneficiary_shares, other_beneficiaries) are emitted
with a header row whose column-0 token identifies the section:
``"Children"``, ``"Successor Trustees"``, ``"Beneficiary Shares"``,
``"Other Beneficiaries"``. The extractor dispatches each table to the
right list-builder by reading the column-0 header. The cycle-4b/5
label/value and checkbox-detection paths still walk every row of every
table; only the list-shaped extractions consult the header.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document  # type: ignore[import-untyped]


def make_docx_with(
    tmp_path: Path,
    *,
    trust_type: str | None = None,  # raw cell text: "Joint" / "Individual"
    marital_status: str | None = None,  # raw cell text: "Married" / "Unmarried"
    grantor_name: str | None = None,
    co_grantor_name: str | None = None,
    children: list[tuple[str, str]] | None = None,  # (name, dob)
    successor_trustees: list[str] | None = None,
    beneficiary_shares: list[tuple[str, str]] | None = None,  # (name, share-percent)
    other_beneficiaries: list[str] | None = None,
    exclusions: str | None = None,  # v2 text-block free text
) -> Path:
    """Construct a minimal ``.docx`` with table rows / paragraphs wired to the
    specified content.

    Each kwarg controls one table or paragraph block per the v2.2
    questionnaire layout. Returns a path under ``tmp_path``. The ``.docx``
    structure mirrors the v2.2 layout closely enough that the parser's
    table-detection logic (ported in cycle 4b, extended in cycles 5 and 6)
    walks it identically.
    """
    out = tmp_path / "fixture.docx"
    doc = Document()

    # Trust-type / marital-status checkbox table. Two-row, two-column shape:
    # col 0 is the check marker ("X"), col 1 is the v2.2 phrasing the parser
    # matches against its checkbox maps. Rows are emitted only when the
    # corresponding kwarg is set, so the cycle-5 "no checkbox" None-gate
    # test exercises an empty fixture cleanly.
    if trust_type is not None or marital_status is not None:
        table = doc.add_table(rows=2, cols=2)
        if trust_type is not None:
            table.cell(0, 0).text = "X"
            table.cell(0, 1).text = f"This is a {trust_type} trust"
        if marital_status is not None:
            table.cell(1, 0).text = "X"
            table.cell(1, 1).text = f"Grantor is {marital_status}"

    # Grantor / co-grantor name rows (2-col label/value form; parser's
    # _LABEL_KEY_MAP recognizes the col-0 label).
    if grantor_name is not None or co_grantor_name is not None:
        name_tbl = doc.add_table(rows=2, cols=2)
        if grantor_name is not None:
            name_tbl.cell(0, 0).text = "Grantor full legal name"
            name_tbl.cell(0, 1).text = grantor_name
        if co_grantor_name is not None:
            name_tbl.cell(1, 0).text = "Co-grantor full legal name"
            name_tbl.cell(1, 1).text = co_grantor_name

    # Children table — header row + (name, dob) per data row.
    if children:
        child_tbl = doc.add_table(rows=len(children) + 1, cols=2)
        child_tbl.cell(0, 0).text = "Children"
        child_tbl.cell(0, 1).text = "Date of Birth"
        for row_idx, (name, dob) in enumerate(children, start=1):
            child_tbl.cell(row_idx, 0).text = name
            child_tbl.cell(row_idx, 1).text = dob

    # Successor-trustees table — header row + single name column per row.
    if successor_trustees:
        tr_tbl = doc.add_table(rows=len(successor_trustees) + 1, cols=1)
        tr_tbl.cell(0, 0).text = "Successor Trustees"
        for row_idx, name in enumerate(successor_trustees, start=1):
            tr_tbl.cell(row_idx, 0).text = name

    # Beneficiary-shares table — header row + (name, share-percent) per row.
    if beneficiary_shares:
        bs_tbl = doc.add_table(rows=len(beneficiary_shares) + 1, cols=2)
        bs_tbl.cell(0, 0).text = "Beneficiary Shares"
        bs_tbl.cell(0, 1).text = "Share %"
        for row_idx, (name, share) in enumerate(beneficiary_shares, start=1):
            bs_tbl.cell(row_idx, 0).text = name
            bs_tbl.cell(row_idx, 1).text = share

    # Other-beneficiaries table — header row + single name column per row.
    if other_beneficiaries:
        ob_tbl = doc.add_table(rows=len(other_beneficiaries) + 1, cols=1)
        ob_tbl.cell(0, 0).text = "Other Beneficiaries"
        for row_idx, name in enumerate(other_beneficiaries, start=1):
            ob_tbl.cell(row_idx, 0).text = name

    # Exclusions paragraph (v2 text-block free text; parsed as the
    # parser-internal exclusions_string carrier per F3 finding). The
    # parser detects the "Exclusions:" prefix and strips it before
    # tokenizing.
    if exclusions is not None:
        doc.add_paragraph(f"Exclusions: {exclusions}")

    doc.save(str(out))
    return out
