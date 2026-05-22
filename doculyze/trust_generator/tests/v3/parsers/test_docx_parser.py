"""Unit tests for parse_docx and the post-promotion protocol helper."""

from __future__ import annotations

import pytest
from docx import Document

from trust_generator.v3.schema import (
    MaritalStatus,
    QuestionnaireSeed,
    TrustType,
    promote_seed,
)


def test_parse_docx_smoke(tmp_path):
    """parse_docx exists, opens a minimal valid .docx, and returns a TrustData.

    The Red signal is unambiguously "parser module absent" — the minimal.docx
    has no v2.2-shape dependencies, so cycle-4b asset-shape failures cannot
    masquerade as cycle-4a failures (plan-review pass 1, M5 cycle-split).

    Import is via the explicit module path (`v3.parsers.docx_parser`) rather
    than the package re-export. The package `__init__.py` re-export of
    `parse_docx` lands in the downstream `registry` sibling's cycle 9; until
    that lands, every test in this plan imports through the explicit module
    path.
    """
    from trust_generator.v3.parsers.docx_parser import (
        parse_docx,  # NOQA: deliberate late import
    )

    minimal_docx = tmp_path / "minimal.docx"
    doc = Document()
    doc.add_paragraph("placeholder")
    doc.save(str(minimal_docx))

    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)
    seed_snapshot = seed_initialized.model_copy(deep=True)

    result = parse_docx(minimal_docx, seed_initialized)

    assert result is not None
    # P3 invariant (spec §4): seed_initialized is field-level equal before and
    # after. The deepcopy at parser entry is the reference implementation that
    # satisfies this postcondition; the test is implementation-agnostic.
    assert seed_initialized == seed_snapshot
    # Deepcopy proof (spec §5.3 step 1): the returned TrustData is a separate
    # instance from the caller-supplied seed_initialized.
    assert result is not seed_initialized


def test_parse_docx_synthetic_grantor_name_extraction(tmp_path):
    """A synthetic fixture with a single Grantor row populates result.grantor.

    Independent of asset availability; pins that the cycle-4b flat-key
    extraction wires at least one v2.2 row into result. The full row
    coverage lives in cycle 6 (coercion integration); cycle 4b's
    obligation is just that *some* table content survives the parse.
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx

    fixture = tmp_path / "single_row.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Grantor full legal name"
    table.cell(0, 1).text = "John Andrew Doe"
    table.cell(1, 0).text = "Grantor date of birth"
    table.cell(1, 1).text = "01/15/1970"
    doc.save(str(fixture))

    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)
    result = parse_docx(fixture, seed_initialized)

    assert result.grantor.full_legal_name == "John Andrew Doe"


# ---------------------------------------------------------------------------
# Cycle 5 — post-promotion contract (spec §6.7)
# ---------------------------------------------------------------------------

from tests.v3.parsers._docx_fixtures import make_docx_with
from trust_generator.v3.schema import GrantorInfo


def test_grantor_info_default_constructor_equality_is_deterministic():
    """Pinned precondition for the chore #37 dematerialization branch.

    The dematerialization branch in `_apply_post_promotion_protocol`
    checks `result.co_grantor == GrantorInfo()` to detect a
    seed-materialized-empty co_grantor. This relies on Pydantic v2
    BaseModel field-equality being deterministic for default-constructed
    GrantorInfo instances.

    A future schema change that introduces a non-deterministic default
    (e.g., `default_factory=uuid.uuid4` or `default_factory=datetime.now`)
    would silently break the branch: `GrantorInfo() != GrantorInfo()` and
    the elif-branch never fires, leaving seed-materialized-empty
    co_grantors in place across (IN, UM) transitions. This single-line
    test catches the regression at the next pixi run test.

    If this test ever fails, fix the branch by switching the equality
    check to `result.co_grantor.model_dump(exclude_defaults=True) == {}`,
    which is robust against non-deterministic defaults.
    """
    assert GrantorInfo() == GrantorInfo()


@pytest.mark.parametrize(
    (
        "seed_state",
        "parsed_state",
        "expected_grantor_caption",
        "expected_co_grantor_caption",
        "expected_co_grantor_present",
    ),
    [
        # (JT, MR) -> (IN, UM): joint mutation; captions collapse; default-only
        # co_grantor dematerializes (chore #37 amendment).
        (
            (TrustType.JOINT, MaritalStatus.MARRIED),
            (TrustType.INDIVIDUAL, MaritalStatus.UNMARRIED),
            "Grantor",
            "Spouse",
            False,
        ),
        # (IN, UM) -> (JT, MR): re-materialization (None → GrantorInfo()).
        (
            (TrustType.INDIVIDUAL, MaritalStatus.UNMARRIED),
            (TrustType.JOINT, MaritalStatus.MARRIED),
            "Grantor A",
            "Grantor B",
            True,
        ),
        # (JT, MR) -> (IN, MR): caption-only mutation; co_grantor preserved
        # (post-state still requires co_grantor via MARRIED).
        (
            (TrustType.JOINT, MaritalStatus.MARRIED),
            (TrustType.INDIVIDUAL, MaritalStatus.MARRIED),
            "Grantor",
            "Spouse",
            True,
        ),
        # (IN, MR) -> (JT, MR): caption mutation across marital-equivalent
        # transition; co_grantor preserved.
        (
            (TrustType.INDIVIDUAL, MaritalStatus.MARRIED),
            (TrustType.JOINT, MaritalStatus.MARRIED),
            "Grantor A",
            "Grantor B",
            True,
        ),
    ],
    ids=["jt_mr_to_in_um", "in_um_to_jt_mr", "jt_mr_to_in_mr", "in_mr_to_jt_mr"],
)
def test_post_promotion_protocol_combinatorial(
    seed_state,
    parsed_state,
    expected_grantor_caption,
    expected_co_grantor_caption,
    expected_co_grantor_present,
    tmp_path,
):
    """All four (seed, parsed) trust_type/marital_status combinations apply
    correctly per spec §5.3 step 4. Combinatorial cycle-5 coverage rule.
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx

    seed_trust_type, seed_marital = seed_state
    parsed_trust_type, parsed_marital = parsed_state

    seed = QuestionnaireSeed(trust_type=seed_trust_type, marital_status=seed_marital)
    seed_initialized = promote_seed(seed)

    fixture = make_docx_with(
        tmp_path,
        trust_type="Joint" if parsed_trust_type == TrustType.JOINT else "Individual",
        marital_status="Married" if parsed_marital == MaritalStatus.MARRIED else "Unmarried",
    )

    result = parse_docx(fixture, seed_initialized)

    assert result.trust_id.trust_type == parsed_trust_type
    assert result.trust_id.marital_status == parsed_marital
    assert result.trust_id.grantor_caption == expected_grantor_caption
    assert result.trust_id.co_grantor_caption == expected_co_grantor_caption
    assert (result.co_grantor is not None) == expected_co_grantor_present


def test_post_promotion_protocol_none_gate_preserves_seed_value(tmp_path):
    """F1 finding (plan-review pass 2): parsed-None means "no mutation
    requested"; the seed-initialized value persists. The None-gate is
    load-bearing because trust_type is a required schema field — assigning
    None would breach Pydantic validation.

    Fixture has no trust_type / marital_status checkboxes; the parser's
    flat-key extraction emits None for both. The seed-initialized state
    (JT+MR captions, materialized co_grantor) is unchanged.
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx

    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)
    fixture = make_docx_with(tmp_path)  # no trust_type / marital_status rows

    result = parse_docx(fixture, seed_initialized)

    assert result.trust_id.trust_type == TrustType.JOINT
    assert result.trust_id.marital_status == MaritalStatus.MARRIED
    assert result.trust_id.grantor_caption == "Grantor A"
    assert result.trust_id.co_grantor_caption == "Grantor B"
    assert result.co_grantor is not None


def test_parser_preserves_populated_co_grantor_on_marital_transition(tmp_path):
    """Already-populated co_grantor survives a marital_status change.

    Spec §5.3 step 4 sub-bullet: "If co_grantor is already populated,
    preserve it (the populated data is meaningful; the mutation decided
    the grantor exists, not what their identity is)."
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx

    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)
    # Populate co_grantor BEFORE parsing — simulates a downstream caller
    # that filled co_grantor identity from a prior workflow.
    seed_initialized.co_grantor = GrantorInfo(full_legal_name="Jane Doe")
    snapshot_co_grantor = seed_initialized.co_grantor.model_copy(deep=True)

    fixture = make_docx_with(
        tmp_path,
        trust_type="Individual",  # IN+MR — still requires co_grantor
        marital_status="Married",
    )

    result = parse_docx(fixture, seed_initialized)

    assert result.trust_id.trust_type == TrustType.INDIVIDUAL
    assert result.trust_id.marital_status == MaritalStatus.MARRIED
    assert result.co_grantor is not None
    assert result.co_grantor.full_legal_name == "Jane Doe"
    # P3 invariant: caller's seed_initialized.co_grantor is field-level
    # equal pre- vs. post-call.
    assert seed_initialized.co_grantor == snapshot_co_grantor


def test_parser_preserves_populated_co_grantor_under_dematerialization_target(
    tmp_path,
):
    """Chore #37 preservation rule: a POPULATED co_grantor survives a
    (JT, MR) → (IN, UM) transition even though the post-mutation state
    requires no co_grantor.

    The dematerialization branch fires only when the current co_grantor
    is field-equal to GrantorInfo() (default-only). Any populated field
    breaks the equality → preservation fallback → co_grantor retained.
    This is the regression guard for the chore #37 amendment's
    "populated data is meaningful" carve-out: parsers must not drop
    paralegal-supplied co_grantor identity when a competing trust-type
    or marital-status mutation would otherwise dematerialize the slot.
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx

    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)
    seed_initialized.co_grantor = GrantorInfo(full_legal_name="Jane Doe")
    snapshot_co_grantor = seed_initialized.co_grantor.model_copy(deep=True)

    fixture = make_docx_with(
        tmp_path,
        trust_type="Individual",
        marital_status="Unmarried",
    )

    result = parse_docx(fixture, seed_initialized)

    # Post-mutation state requires no co_grantor, but the populated
    # identity is preserved by the chore #37 carve-out.
    assert result.trust_id.trust_type == TrustType.INDIVIDUAL
    assert result.trust_id.marital_status == MaritalStatus.UNMARRIED
    assert result.co_grantor is not None
    assert result.co_grantor.full_legal_name == "Jane Doe"
    # P3 invariant: caller's seed_initialized.co_grantor field-level equal
    # pre- vs. post-call.
    assert seed_initialized.co_grantor == snapshot_co_grantor


def test_parser_never_reinvokes_promote_seed(tmp_path):
    """Spec §5.3 step 4 ("parsers do not call promote_seed under any branch")
    and §4 P1 invariant ("parsers never re-invoke promote_seed").
    """
    from unittest.mock import patch

    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)
    fixture = make_docx_with(
        tmp_path,
        trust_type="Individual",
        marital_status="Unmarried",
    )

    with patch(
        "trust_generator.v3.parsers.docx_parser.promote_seed"
    ) as mock_promote:
        from trust_generator.v3.parsers.docx_parser import parse_docx
        parse_docx(fixture, seed_initialized)
        assert mock_promote.call_count == 0


# ---------------------------------------------------------------------------
# Cycle 6 — coercion integration + post-merge resolution (spec §6.8)
# ---------------------------------------------------------------------------

import logging


def test_coercion_integration_malformed_date_falls_back_to_None(tmp_path, caplog):
    """A malformed date in a child DOB cell coerces to None and emits a
    log.warning. The parse succeeds (no exception)."""
    from trust_generator.v3.parsers.docx_parser import parse_docx

    fixture = make_docx_with(
        tmp_path,
        grantor_name="John Andrew Doe",
        children=[("Alice Doe", "sometime in 2010")],  # unparseable date
    )
    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)

    with caplog.at_level(logging.WARNING, logger="trust_generator.v3.parsers"):
        result = parse_docx(fixture, seed_initialized)

    assert result.grantor.full_legal_name == "John Andrew Doe"
    assert len(result.children) == 1
    assert result.children[0].full_legal_name == "Alice Doe"
    assert result.children[0].date_of_birth is None
    assert any("could not parse date" in rec.message.lower() for rec in caplog.records)


def test_coercion_integration_one_token_corporate_trustee_routes_to_CorporateTrustee(
    tmp_path, caplog
):
    """A trustee name matching the §5.4.9 entity heuristic routes to
    CorporateTrustee with an INFO log. (Name is "First National Bank" —
    multi-token; the test name's "one_token" is historical, referring to
    the §5.4.4 trap; the actual mechanism here is §5.4.9's suffix
    heuristic re-applied inside `_apply_post_merge_resolution`.)
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx
    from trust_generator.v3.schema import CorporateTrustee

    fixture = make_docx_with(
        tmp_path,
        successor_trustees=["First National Bank"],
    )
    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)

    with caplog.at_level(logging.INFO, logger="trust_generator.v3.parsers"):
        result = parse_docx(fixture, seed_initialized)

    assert len(result.successor_trustees) == 1
    assert isinstance(result.successor_trustees[0], CorporateTrustee)
    assert result.successor_trustees[0].entity_name == "First National Bank"
    assert any(
        "discriminated as CorporateTrustee" in rec.message
        or "CorporateTrustee" in rec.message
        for rec in caplog.records
    )


def test_coercion_integration_placeholder_prefix_stripped_from_person_reference(
    tmp_path,
):
    """A v2-corpus cell typed beside a bracketed hint (per spec §5.4.4 +
    Decision log #12) has the prefix stripped before _to_person_reference."""
    from trust_generator.v3.parsers.docx_parser import parse_docx

    fixture = make_docx_with(
        tmp_path,
        co_grantor_name="[Spouse's full legal name] Jane Doe",
    )
    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)

    result = parse_docx(fixture, seed_initialized)

    assert result.co_grantor is not None
    assert result.co_grantor.full_legal_name == "Jane Doe"


def test_coercion_integration_disinherit_multi_match_warns_and_picks_iteration_order(
    tmp_path, caplog
):
    """F4 finding (plan-review pass 2): a v2 exclusions token matching
    multiple beneficiaries across the fixed iteration order
    (children → descendants → other_beneficiaries) emits exactly one
    disinherit flip on the iteration-order-first match plus one WARNING
    naming both candidates.

    Token: "John". Match candidates: "John Smith" (children),
    "Johnny Doe" (other_beneficiaries). Expected: only "John Smith" is
    flipped to disinherit=True; the other_beneficiaries entry is
    preserved with disinherit=False; one WARNING naming both names.
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx

    fixture = make_docx_with(
        tmp_path,
        children=[("John Smith", "2010-01-01")],
        other_beneficiaries=["Johnny Doe"],
        exclusions="John",
    )
    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)

    with caplog.at_level(logging.WARNING, logger="trust_generator.v3.parsers"):
        result = parse_docx(fixture, seed_initialized)

    # F2 iteration-order pin (plan-review pass 2): children first.
    assert len(result.children) == 1
    assert result.children[0].full_legal_name == "John Smith"
    assert result.children[0].disinherit is True
    assert result.children[0].disinherit_reason == "John"

    # The other_beneficiaries entry retains disinherit=False.
    assert len(result.other_beneficiaries) == 1
    assert result.other_beneficiaries[0].full_legal_name == "Johnny Doe"
    assert result.other_beneficiaries[0].disinherit is False

    # Exactly one WARNING naming both candidates.
    multi_match_warnings = [
        rec
        for rec in caplog.records
        if rec.levelno == logging.WARNING
        and "John Smith" in rec.message
        and "Johnny Doe" in rec.message
    ]
    assert len(multi_match_warnings) == 1


def test_coercion_integration_unmatched_exclusion_token_flows_to_external_exclusions(
    tmp_path,
):
    """Spec §5.4.10 algorithm step 4: an exclusions token with no beneficiary
    match flows to result.external_exclusions as a PersonReference, with
    result.external_exclusion_reasons[token] = token."""
    from trust_generator.v3.parsers.docx_parser import parse_docx

    fixture = make_docx_with(
        tmp_path,
        children=[("Alice Doe", "2010-01-01")],
        exclusions="Bob Roe",  # no match in children/descendants/other_beneficiaries
    )
    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)

    result = parse_docx(fixture, seed_initialized)

    assert len(result.external_exclusions) == 1
    assert result.external_exclusions[0].full_legal_name == "Bob Roe"
    assert result.external_exclusion_reasons.get("Bob Roe") == "Bob Roe"
    # The unmatched-only case must NOT flip any beneficiary's disinherit.
    assert result.children[0].disinherit is False


def test_coercion_integration_unparseable_share_percent_drops_row(
    tmp_path, caplog
):
    """Unparseable share-percent ("a lot") → the beneficiary_shares row is
    dropped per the §5.4.2 share-percent branch (Decision log #11), and a
    log.warning is emitted.
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx

    fixture = make_docx_with(
        tmp_path,
        beneficiary_shares=[("Alice Doe", "a lot")],
    )
    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)

    with caplog.at_level(logging.WARNING, logger="trust_generator.v3.parsers"):
        result = parse_docx(fixture, seed_initialized)

    # Share-percent row should be DROPPED per spec §5.4.2 share-percent
    # branch (Decision log #11), NOT coerced to Decimal(0). Assert empty.
    assert result.beneficiary_shares == []
    assert any("could not parse" in rec.message.lower() for rec in caplog.records)


def test_co_grantor_one_token_name_preserves_data_as_entity(tmp_path):
    """C2 (PR #12 review): a one-token co_grantor cell must not wipe the name.

    ``_to_person_reference`` returns entity-form for a one-token cell
    (``full_legal_name=""``, ``is_entity=True``, ``entity_name=<cell>``).
    The docx co_grantor path must copy all three PersonReference fields —
    mirroring the children / successor_trustees / other_beneficiaries paths
    in the same function — so the data survives in ``entity_name`` rather
    than the name being silently overwritten with an empty string.
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx

    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT, marital_status=MaritalStatus.MARRIED
    )
    seed_initialized = promote_seed(seed)  # JT+MR → co_grantor materialized
    fixture = make_docx_with(tmp_path, co_grantor_name="Smith")
    result = parse_docx(fixture, seed_initialized)

    assert result.co_grantor is not None
    assert result.co_grantor.is_entity is True
    assert result.co_grantor.entity_name == "Smith"
    assert result.co_grantor.full_legal_name == ""


# ---------------------------------------------------------------------------
# Chore #48 — no-handler diagnostic for multi-row unrecognised tables
# ---------------------------------------------------------------------------


def test_extract_flat_emits_warning_for_multi_row_table_matching_no_handler(
    tmp_path, caplog
):
    """A multi-row table whose rows match neither a list header, a checkbox,
    nor a label/value key produces a WARNING so template drift is visible.

    The fixture constructs a 3-row, 2-col table where column 0 is an
    unknown token — representative of a questionnaire section added in a
    template revision that the parser has not yet been updated to recognise.
    The parse must NOT raise; the only change is the warning emission.
    """
    from trust_generator.v3.parsers.docx_parser import parse_docx

    fixture = tmp_path / "unrecognised_table.docx"
    doc = Document()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "Unknown Section Header"
    tbl.cell(0, 1).text = ""
    tbl.cell(1, 0).text = "Some value"
    tbl.cell(1, 1).text = "Data"
    tbl.cell(2, 0).text = "Another value"
    tbl.cell(2, 1).text = "More data"
    doc.save(str(fixture))

    seed = QuestionnaireSeed(
        trust_type=TrustType.JOINT,
        marital_status=MaritalStatus.MARRIED,
    )
    seed_initialized = promote_seed(seed)

    with caplog.at_level(logging.WARNING, logger="trust_generator.v3.parsers"):
        result = parse_docx(fixture, seed_initialized)

    # Parse completes without error; no field is populated from the unknown table.
    assert result is not None
    # A WARNING naming the unrecognised first-cell text must appear.
    unrecognised_warnings = [
        rec
        for rec in caplog.records
        if rec.levelno == logging.WARNING and "Unknown Section Header" in rec.message
    ]
    assert len(unrecognised_warnings) == 1
